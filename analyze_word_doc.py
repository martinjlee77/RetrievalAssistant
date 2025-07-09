#!/usr/bin/env python3
"""
Analyze Word document structure and content for RAG suitability
"""
import os
from docx import Document
from docx.shared import Inches
import json

def analyze_word_document(file_path):
    """Analyze Word document structure and content"""
    
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return
    
    try:
        # Load the document
        doc = Document(file_path)
        
        analysis = {
            'document_info': {
                'total_paragraphs': len(doc.paragraphs),
                'total_tables': len(doc.tables),
                'sections': len(doc.sections)
            },
            'content_structure': [],
            'tables_found': [],
            'text_sample': [],
            'headings_hierarchy': [],
            'potential_removals': []
        }
        
        # Analyze paragraphs and structure
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
                
            # Check for headings
            if paragraph.style.name.startswith('Heading'):
                analysis['headings_hierarchy'].append({
                    'level': paragraph.style.name,
                    'text': text,
                    'paragraph_index': i
                })
            
            # Sample first 50 paragraphs for structure analysis
            if i < 50:
                analysis['content_structure'].append({
                    'index': i,
                    'style': paragraph.style.name,
                    'text': text[:200] + '...' if len(text) > 200 else text,
                    'length': len(text)
                })
            
            # Collect text samples
            if len(text) > 50 and len(analysis['text_sample']) < 20:
                analysis['text_sample'].append({
                    'paragraph': i,
                    'text': text[:300] + '...' if len(text) > 300 else text
                })
        
        # Analyze tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            analysis['tables_found'].append({
                'table_index': i,
                'rows': len(table.rows),
                'columns': len(table.rows[0].cells) if table.rows else 0,
                'sample_data': table_data[:3]  # First 3 rows
            })
        
        # Identify potential content for removal
        removal_candidates = []
        
        for para in analysis['content_structure']:
            text = para['text'].lower()
            
            # Headers/footers
            if any(keyword in text for keyword in ['page', 'chapter', 'ey building', 'copyright']):
                removal_candidates.append(f"Header/Footer: {para['text'][:100]}")
            
            # Navigation elements
            if any(keyword in text for keyword in ['table of contents', 'index', 'appendix']):
                removal_candidates.append(f"Navigation: {para['text'][:100]}")
            
            # Metadata
            if any(keyword in text for keyword in ['publication date', 'version', 'document id']):
                removal_candidates.append(f"Metadata: {para['text'][:100]}")
        
        analysis['potential_removals'] = list(set(removal_candidates))
        
        return analysis
        
    except Exception as e:
        print(f"Error analyzing document: {str(e)}")
        return None

def print_analysis(analysis):
    """Print formatted analysis results"""
    
    print("=== WORD DOCUMENT ANALYSIS ===\n")
    
    print("📊 DOCUMENT OVERVIEW:")
    info = analysis['document_info']
    print(f"  • Total paragraphs: {info['total_paragraphs']}")
    print(f"  • Total tables: {info['total_tables']}")
    print(f"  • Sections: {info['sections']}")
    print()
    
    print("📋 HEADING HIERARCHY:")
    for heading in analysis['headings_hierarchy'][:10]:  # Show first 10
        print(f"  • {heading['level']}: {heading['text']}")
    print()
    
    print("📄 CONTENT STRUCTURE SAMPLE:")
    for item in analysis['content_structure'][:15]:  # Show first 15
        print(f"  [{item['index']}] {item['style']}: {item['text'][:100]}...")
    print()
    
    print("📊 TABLES FOUND:")
    for table in analysis['tables_found'][:5]:  # Show first 5
        print(f"  • Table {table['table_index']}: {table['rows']}x{table['columns']} cells")
        if table['sample_data']:
            print(f"    Sample: {table['sample_data'][0]}")
    print()
    
    print("🗑️  POTENTIAL REMOVALS:")
    for removal in analysis['potential_removals'][:10]:
        print(f"  • {removal}")
    print()
    
    print("📝 TEXT SAMPLES:")
    for sample in analysis['text_sample'][:5]:
        print(f"  [{sample['paragraph']}] {sample['text']}")
        print()

if __name__ == "__main__":
    file_path = "attached_assets/ey-frdbb3043-09-24-2024_1752054352087.docx"
    
    print("Analyzing Word document for RAG suitability...")
    print(f"File: {file_path}")
    print()
    
    analysis = analyze_word_document(file_path)
    
    if analysis:
        print_analysis(analysis)
        
        # RAG Suitability Assessment
        print("🎯 RAG SUITABILITY ASSESSMENT:")
        
        total_content = analysis['document_info']['total_paragraphs']
        tables_count = analysis['document_info']['total_tables']
        headings_count = len(analysis['headings_hierarchy'])
        
        score = 0
        
        # Content volume
        if total_content > 200:
            score += 25
            print("  ✅ Excellent content volume")
        elif total_content > 100:
            score += 15
            print("  ⚠️  Moderate content volume")
        else:
            score += 5
            print("  ❌ Low content volume")
        
        # Structure quality
        if headings_count > 10:
            score += 25
            print("  ✅ Well-structured with clear headings")
        elif headings_count > 5:
            score += 15
            print("  ⚠️  Some structure present")
        else:
            score += 5
            print("  ❌ Poor structure")
        
        # Table content
        if tables_count > 5:
            score += 25
            print("  ✅ Rich table content")
        elif tables_count > 2:
            score += 15
            print("  ⚠️  Some tables present")
        else:
            score += 5
            print("  ❌ Limited table content")
        
        # Content quality
        avg_paragraph_length = sum(len(p['text']) for p in analysis['content_structure']) / len(analysis['content_structure'])
        if avg_paragraph_length > 100:
            score += 25
            print("  ✅ Substantial paragraph content")
        elif avg_paragraph_length > 50:
            score += 15
            print("  ⚠️  Moderate paragraph content")
        else:
            score += 5
            print("  ❌ Short paragraph content")
        
        print(f"\n📈 OVERALL RAG SUITABILITY: {score}/100")
        
        if score >= 80:
            print("  🎉 EXCELLENT - Highly suitable for RAG system")
        elif score >= 60:
            print("  ✅ GOOD - Suitable for RAG with minor cleanup")
        elif score >= 40:
            print("  ⚠️  MODERATE - Requires significant cleanup")
        else:
            print("  ❌ POOR - Not suitable for RAG in current form")
    else:
        print("❌ Failed to analyze document")