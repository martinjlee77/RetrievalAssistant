"""
Clean Memo Generator - No Text Processing
Displays clean GPT-4o output without any corruption.
"""

import streamlit as st
import logging
from datetime import datetime
from typing import Dict, Any
from shared.disclaimer_generator import DisclaimerGenerator
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import tempfile
import base64


logger = logging.getLogger(__name__)


class CleanMemoGenerator:
    """Clean memo generator that preserves GPT-4o text exactly as generated."""

    def __init__(self, template_path=None):
        """Initialize - template_path ignored for now."""
        pass

    def _clean_po_summary_blocks(self, text: str) -> str:
        """Remove [BEGIN_PO_SUMMARY]...[END_PO_SUMMARY] blocks from analysis text."""
        import re
        # Remove PO summary blocks that are meant for executive summary generation only
        pattern = r'\[BEGIN_PO_SUMMARY\].*?\[END_PO_SUMMARY\]'
        cleaned_text = re.sub(pattern, '', text, flags=re.DOTALL)
        # Clean up any extra whitespace left behind
        cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)
        return cleaned_text.strip()

    def combine_clean_steps(self,
                            analysis_results: Dict[str, Any],
                            analysis_id: str | None = None) -> str:
        """Combine clean step markdown into final memo - NO PROCESSING."""

        # Get basic info
        customer_name = analysis_results.get('customer_name', 'Customer')
        analysis_title = analysis_results.get('analysis_title',
                                              'Contract Analysis')
        analysis_date = datetime.now().strftime("%B %d, %Y")

        # Build memo with memo ID and disclaimer at very top
        memo_lines = []

        # Add memo ID at the very top if provided
        if analysis_id:
            memo_lines.extend([
                f"**MEMO ID:** {analysis_id}",
                "",
            ])

        # Add disclaimer
        memo_lines.extend([
            DisclaimerGenerator.get_top_banner(), "", "# ASC 606 MEMORANDUM",
            "", f"**TO:** Chief Accounting Officer",
            f"**FROM:** Technical Accounting Team - AI",
            f"**DATE:** {analysis_date}",
            f"**RE:** {analysis_title} - ASC 606 Revenue Recognition Analysis",
            f"**DOCUMENTS REVIEWED:** {analysis_results.get('filename', 'Contract Documents')}",
            "", ""
        ])

        # Add Executive Summary
        if 'executive_summary' in analysis_results:
            memo_lines.extend([
                "## EXECUTIVE SUMMARY", "",
                analysis_results['executive_summary'], "", ""
            ])

        # Add Background
        memo_lines.extend(["## BACKGROUND", ""])

        if 'background' in analysis_results:
            memo_lines.extend([analysis_results['background'], "", ""])
        else:
            memo_lines.extend([
                f"We have reviewed the contract documents provided by {customer_name} to determine the appropriate revenue recognition treatment under ASC 606. This memorandum presents our analysis following the five-step ASC 606 methodology.",
                "", ""
            ])

        # Add Analysis Section Header
        memo_lines.extend([
            "## ASC 606 ANALYSIS",
        ])

        # Add each step's clean markdown content - check both locations
        steps_added = 0
        for step_num in range(1, 6):
            step_key = f'step_{step_num}'
            step_data = None

            # Check if steps are in analysis_results directly
            if step_key in analysis_results:
                step_data = analysis_results[step_key]
            # Check if steps are in analysis_results['steps']
            elif 'steps' in analysis_results and step_key in analysis_results[
                    'steps']:
                step_data = analysis_results['steps'][step_key]

            if step_data and isinstance(
                    step_data, dict) and 'markdown_content' in step_data:
                # Add clean content directly - ZERO PROCESSING (but remove PO summary blocks)
                clean_content = step_data['markdown_content']
                clean_content = self._clean_po_summary_blocks(clean_content)
                memo_lines.append(clean_content)
                memo_lines.append("")
                steps_added += 1
                logger.info(
                    f"Added clean step {step_num} content ({len(clean_content)} chars)"
                )
            else:
                logger.warning(
                    f"Step {step_num} not found or missing markdown_content. Available keys: {list(analysis_results.keys())}"
                )

        # Add Conclusion Section
        if 'conclusion' in analysis_results:
            memo_lines.extend([
                "",
                "## CONCLUSION",
                "",
                analysis_results['conclusion'],
                "",
            ])

        # Add footer with full disclaimer
        memo_lines.extend([
            "---", "", "**PREPARED BY:** [Analyst Name] | [Title] | [Date]",
            "**REVIEWED BY:** [Reviewer Name] | [Title] | [Date]", "",
            DisclaimerGenerator.get_full_disclaimer()
        ])

        # Join and return - NO PROCESSING
        final_memo = "\n".join(memo_lines)
        logger.info(
            f"Clean memo generated: {len(final_memo)} chars, {steps_added}/5 steps"
        )
        return final_memo

    def _generate_pdf(self, memo_content: str) -> bytes | None:
        """Generate PDF using ReportLab"""
        try:
            from shared.pdf_generator import generate_pdf_from_markdown
            
            logger.info("Generating PDF using ReportLab")
            pdf_bytes = generate_pdf_from_markdown(memo_content)
            
            if pdf_bytes:
                logger.info(f"PDF generation successful: {len(pdf_bytes)} bytes")
                return pdf_bytes
            else:
                logger.error("PDF generation returned None")
                return None
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return None

    def _clean_html_tags(self, text: str) -> str:
        """Remove HTML tags from text for clean DOCX output."""
        import re
        # Remove ALL HTML tags for clean Word output
        text = re.sub(r'<[^>]+>', '', text)  # Remove any HTML tag
        # Convert common HTML entities
        text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        return text.strip()

    def _convert_markdown_to_html(self, markdown_content: str) -> str:
        """Convert memo markdown to clean HTML for PDF generation."""
        lines = markdown_content.split('\n')
        html_lines = []
        in_list = False
        
        for line in lines:
            # Headers
            if line.startswith('# '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h2 style="color: #ffffff;">{line[3:]}</h2>')
            elif line.startswith('### '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h3 style="color: #ffffff;">{line[4:]}</h3>')
            # Lists
            elif line.strip().startswith('- '):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                html_lines.append(f'<li>{line.strip()[2:]}</li>')
            # Tables (simple conversion)
            elif '|' in line and line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                # Simple table handling
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    if line.strip().startswith('|'):
                        row_html = '<tr>' + ''.join(f'<td>{cell}</td>' for cell in cells) + '</tr>'
                        html_lines.append(row_html)
            # Bold text (enhanced handling for multiple ** pairs)
            elif '**' in line and line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                # Handle multiple bold sections
                processed_line = line
                while '**' in processed_line:
                    processed_line = processed_line.replace('**', '<strong>', 1).replace('**', '</strong>', 1)
                html_lines.append(f'<p style="margin: 10px 0; line-height: 1.6;">{processed_line}</p>')
            # Skip horizontal rules (as requested)
            elif line.strip() == '---':
                continue
            # Empty lines - add small spacing
            elif line.strip() == '':
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                continue
            # Regular paragraphs
            elif line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<p style="margin: 10px 0; line-height: 1.6;">{line}</p>')

        # Close any open lists
        if in_list:
            html_lines.append('</ul>')

        # Join with improved styling - Original fonts for HTML, VLSerif only for PDF italics
        html_content = f"""
        <div style="font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif; 
        line-height: 1.6; max-width: 850px; padding: 30px; 
        border: 1px solid #e1e5e9; 
        border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            {''.join(html_lines)}
        </div>
        """

        return html_content

    def _generate_docx(self, memo_content: str) -> bytes | None:
        """Generate DOCX from memo content using python-docx."""
        try:
            doc = Document()

            # Page set up
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Set default font for the whole document
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)
                        
            # Clean HTML tags from content first
            clean_content = self._clean_html_tags(memo_content)

            # Process content line by line
            lines = clean_content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Headers
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.runs[0].font.size = Pt(24)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # Better section spacing for main title
                    heading.paragraph_format.space_before = Pt(18)
                    heading.paragraph_format.space_after = Pt(12)
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    heading.runs[0].font.size = Pt(18)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # C. Better section spacing for section headers
                    heading.paragraph_format.space_before = Pt(15)
                    heading.paragraph_format.space_after = Pt(9)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # C. Better section spacing for step headers
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(6)
                # Bold text
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line[2:-2]).bold = True
                    # B. Add line spacing for bold paragraphs
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                # Bullet points
                elif line.startswith('- '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    # Add line spacing for bullet points
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                # Regular paragraphs
                else:
                    # Remove markdown formatting
                    clean_line = line.replace('**', '').replace('*', '')
                    p = doc.add_paragraph(clean_line)
                    # Add line spacing for regular paragraphs  
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                
            # Save to bytes
            with tempfile.NamedTemporaryFile() as tmp_file:
                doc.save(tmp_file.name)
                tmp_file.seek(0)
                return tmp_file.read()

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return None

    def display_clean_memo(self,
                           memo_content: str,
                           analysis_id: str | None = None,
                           filename: str | None = None,
                           customer_name: str | None = None) -> None:
        """Display clean memo content with enhanced download options."""

        # STRATEGIC FIX: Store analysis_id as memo_id for PDF generation
        self.memo_id = analysis_id or "memo_unknown"
        
        # Validate memo content
        if not memo_content or memo_content.strip() == "":
            st.error("Memo content is empty. Please regenerate the analysis.")
            return

        # Log what we're about to display
        logger.info(
            f"Displaying clean memo sample: {repr(memo_content[:150])}")

        # Convert markdown to HTML manually to bypass Streamlit's markdown processor
        html_content = self._convert_markdown_to_html(memo_content)

        # Use HTML display which preserves formatting
        st.markdown(html_content, unsafe_allow_html=True)

        # Enhanced Download Section - AFTER memo display (for stability)
        if memo_content and len(memo_content.strip()) > 10:
            st.markdown("---")
            st.markdown("### 💾 Save Your Memo")
            st.info(
                "**IMPORTANT:** Choose your preferred format to save this memo before navigating away."
            )

            # Create columns for download buttons
            col1, col2, col3, col4 = st.columns(4)

            # Generate timestamp for consistent filenames
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = f"asc606_memo_{timestamp}"

            with col1:
                # DOCX Download
                docx_data = self._generate_docx(memo_content)
                if docx_data:
                    st.download_button(
                        label="📄 Word (.docx)",
                        data=docx_data,
                        file_name=f"{base_filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{hash(memo_content[:100])}",
                        use_container_width=True)
                else:
                    st.button("📄 Word",
                              disabled=True,
                              use_container_width=True,
                              help="Word generation failed")

            with col2:
                # Markdown Download (Raw Text)
                st.download_button(
                    label="📝 Markdown",
                    data=memo_content.encode('utf-8'),
                    file_name=f"{base_filename}.md",
                    mime="text/markdown",
                    key=f"download_md_{hash(memo_content[:100])}",
                    use_container_width=True)

            with col3:
                # PDF Download
                pdf_data = self._generate_pdf(memo_content)
                if pdf_data:
                    st.download_button(
                        label="📄 PDF",
                        data=pdf_data,
                        file_name=f"{base_filename}.pdf",
                        mime="application/pdf",
                        key=f"download_pdf_{hash(memo_content[:100])}",
                        use_container_width=True)
                else:
                    st.button("📄 PDF",
                              disabled=True,
                              use_container_width=True,
                              help="PDF generation failed")
    
            with col4:
                # One-click copy button with JavaScript
                copy_key = f"copy_{hash(memo_content[:100])}"
                if st.button("📋 Copy to Clipboard",
                             use_container_width=True,
                             key=copy_key):
                    # Escape content for JavaScript
                    escaped_content = memo_content.replace('`', '\\`').replace('$', '\\$').replace('\\', '\\\\').replace('"', '\\"')

                    # Create JavaScript component to copy to clipboard
                    copy_js = f"""
                    <script>
                        function copyToClipboard() {{
                            const textToCopy = `{escaped_content}`;
                            navigator.clipboard.writeText(textToCopy).then(function() {{
                                alert('Memo copied to clipboard!');
                            }}).catch(function(err) {{
                                // Fallback for older browsers
                                const textArea = document.createElement("textarea");
                                textArea.value = textToCopy;
                                document.body.appendChild(textArea);
                                textArea.select();
                                document.execCommand('copy');
                                document.body.removeChild(textArea);
                                alert('Memo copied to clipboard!');
                            }});
                        }}
                        copyToClipboard();
                    </script>
                    """
                    try:
                        import streamlit.components.v1 as components
                        components.html(copy_js, height=0)
                    except:
                        st.success("Content ready to copy manually")
