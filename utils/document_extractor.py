"""
Document Text Extraction Utility
Handles PDF and Word document text extraction for contract analysis
"""

import io
import logging
import re
from typing import Optional, Dict, Any, List
import PyPDF2
import pdfplumber
import docx
from docx import Document
try:
    import fitz  # PyMuPDF - preferred for text extraction
except ImportError:
    fitz = None

class DocumentExtractor:
    """Extract text from various document formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text(self, uploaded_file) -> Dict[str, Any]:
        """
        Extract text from uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Dict containing extracted text and metadata
        """
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        try:
            # Check file size first - enforce 50MB limit
            uploaded_file.seek(0)  # Reset file pointer
            file_content = uploaded_file.read()
            file_size_mb = len(file_content) / (1024 * 1024)  # Convert to MB
            
            # Note: Removed minimum file size check - let scanned PDF detection handle edge cases
            
            # Note: Removed file size limits for enterprise customers
            
            # Reset file pointer for actual extraction
            uploaded_file.seek(0)
            
            if file_extension == 'pdf':
                extraction_result = self._extract_pdf_text(uploaded_file)
            elif file_extension in ['docx']:
                extraction_result = self._extract_word_text(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Check if extraction yielded meaningful text
            extracted_text = extraction_result.get('text', '').strip()
            # Note: Removed early error for short text - let scanned PDF detection handle this case

            # Enhanced word counting with proper tokenization
            word_count = self._count_words(extracted_text)
            extraction_result['word_count'] = word_count
            extraction_result['file_size_mb'] = round(file_size_mb, 2)
            extraction_result['filename'] = uploaded_file.name
            
            # Add page estimate for user context (≈300 words/page)
            extraction_result['estimated_pages'] = max(1, round(word_count / 300))

            return extraction_result
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
            return {
                'text': '',
                'filename': uploaded_file.name,
                'quality_state': 'blocked',
                'error': str(e),
                'pages': 0,
                'word_count': 0,
                'extraction_method': 'error',
                'file_size_mb': 0,
                'estimated_pages': 0
            }
    
    def _count_words(self, text: str) -> int:
        """
        Enhanced word counting with proper tokenization
        Numbers count as words, whitespace/punctuation split
        """
        if not text or not text.strip():
            return 0
        
        # Split by whitespace and punctuation, filter empty strings
        words = re.findall(r'\b\w+\b', text)
        return len(words)
    
    def _extract_pdf_text(self, uploaded_file) -> Dict[str, Any]:
        """Extract text from PDF file using multiple methods"""
        text = ""
        pages = 0
        extraction_method = "none"
        
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            pdf_bytes = uploaded_file.read()
            uploaded_file.seek(0)  # Reset file pointer
            
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                pages = len(pdf.pages)
                text_parts = []
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    extraction_method = "pdfplumber"
                    
        except Exception as e:
            self.logger.warning(f"pdfplumber failed: {str(e)}, trying PyPDF2")
            
        # Method 2: Fallback to PyPDF2 if pdfplumber fails
        if not text.strip():
            try:
                uploaded_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                pages = len(pdf_reader.pages)
                text_parts = []
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    extraction_method = "PyPDF2"
                    
            except Exception as e:
                self.logger.error(f"PyPDF2 also failed: {str(e)}")
                text = ""
        
        # Enhanced scanned PDF detection with detailed analysis on RAW text (before cleaning)
        detection_analysis = self._analyze_text_quality(text, pages)
        
        # Clean up the text AFTER quality analysis
        text = self._clean_text(text)
        is_likely_scanned = detection_analysis['is_likely_scanned']
        
        # If scanned PDF detected, return specific error message with reasons
        if is_likely_scanned:
            return {
                'text': text,
                'pages': pages,
                'word_count': self._count_words(text) if text else 0,
                'extraction_method': extraction_method,
                'is_likely_scanned': True,
                'quality_state': detection_analysis.get('quality_state', 'blocked'),
                'error': 'scanned_pdf_detected',
                'user_message': self._get_scanned_pdf_message(detection_analysis.get('reasons', [])),
                'detection_reasons': detection_analysis.get('reasons', ['Detection failed']),
                'detection_metrics': detection_analysis.get('metrics', {})
            }
        
        return {
            'text': text,
            'pages': pages,
            'word_count': self._count_words(text) if text else 0,
            'extraction_method': extraction_method,
            'is_likely_scanned': False,
            'error': None if text else "No text could be extracted from PDF",
            'quality_state': detection_analysis.get('quality_state', 'good'),
            'detection_reasons': detection_analysis.get('reasons', []),
            'detection_metrics': detection_analysis.get('metrics', {})
        }
    
    def _extract_word_text(self, uploaded_file) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            doc = Document(uploaded_file)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n\n".join(text_parts)
            text = self._clean_text(text)
            
            return {
                'text': text,
                'pages': len(doc.paragraphs) // 50,  # Rough estimate
                'word_count': self._count_words(text) if text else 0,
                'extraction_method': 'python-docx',
                'is_likely_scanned': False,
                'error': None,
                'quality_state': 'good',  # Word docs are typically clean
                'detection_reasons': [],
                'detection_metrics': {}
            }
            
        except Exception as e:
            self.logger.error(f"Word extraction failed: {str(e)}")
            return {
                'text': '',
                'pages': 0,
                'word_count': 0,
                'extraction_method': 'error',
                'error': str(e)
            }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Only keep non-empty lines
                cleaned_lines.append(line)
        
        # Join lines and normalize spacing
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive newlines
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    def validate_extraction(self, extraction_result: Dict[str, Any]) -> Dict[str, Any]:
        """Validate the quality of text extraction"""
        validation_result = {
            'is_valid': True,
            'quality_score': 100,
            'issues': [],
            'recommendations': []
        }
        
        text = extraction_result.get('text', '')
        word_count = extraction_result.get('word_count', 0)
        error = extraction_result.get('error')
        
        # Check for extraction errors
        if error:
            validation_result['is_valid'] = False
            validation_result['quality_score'] = 0
            validation_result['issues'].append(f"Extraction error: {error}")
            validation_result['recommendations'].append("Try converting the document to a different format")
            return validation_result
        
        # Check for minimum text length
        if word_count < 100:
            validation_result['quality_score'] -= 50
            validation_result['issues'].append("Document appears to contain very little text")
            validation_result['recommendations'].append("Verify the document contains readable text content")
        
        # Check for garbled text (common with scanned PDFs)
        if text and self._detect_garbled_text(text):
            validation_result['quality_score'] -= 30
            validation_result['issues'].append("Text may be garbled or from scanned document")
            validation_result['recommendations'].append("Consider using OCR software if document is scanned")
        
        # Check for reasonable contract content
        contract_keywords = ['agreement', 'contract', 'party', 'parties', 'terms', 'conditions', 'payment', 'service', 'goods']
        found_keywords = sum(1 for keyword in contract_keywords if keyword.lower() in text.lower())
        
        if found_keywords < 3:
            validation_result['quality_score'] -= 20
            validation_result['issues'].append("Document may not contain typical contract language")
            validation_result['recommendations'].append("Verify this is a contract document suitable for ASC 606 analysis")
        
        # Final validation
        if validation_result['quality_score'] < 50:
            validation_result['is_valid'] = False
        
        return validation_result
    
    def _detect_garbled_text(self, text: str) -> bool:
        """Detect if text appears to be garbled (e.g., from poor OCR)"""
        if not text:
            return False
        
        # Check for excessive special characters
        special_char_count = sum(1 for char in text if not char.isalnum() and not char.isspace())
        total_chars = len(text)
        
        if total_chars > 0 and (special_char_count / total_chars) > 0.3:
            return True
        
        # Check for very short words (common in garbled text)
        words = text.split()
        if len(words) > 10:
            very_short_words = sum(1 for word in words if len(word) <= 2)
            if (very_short_words / len(words)) > 0.5:
                return True
        
        return False
    
    def _detect_scanned_pdf(self, pages: int, extracted_text: str) -> bool:
        """Detect if PDF is likely scanned/image-based using strict quality-based metrics"""
        if not extracted_text or not extracted_text.strip():
            return True  # No text extracted at all
        
        # Analyze raw text quality (before cleaning) - this is key for catching OCR artifacts
        detection_result = self._analyze_text_quality(extracted_text, pages)
        
        # Return True if ANY strict quality threshold is exceeded (defensive approach)
        return detection_result['is_likely_scanned']
    
    def _analyze_text_quality(self, raw_text: str, pages: int) -> dict:
        """Comprehensive text quality analysis for scanned PDF detection"""
        import re
        
        text = raw_text.strip()
        if not text:
            return {
                'is_likely_scanned': True, 
                'quality_state': 'blocked',
                'reasons': ['No text extracted'],
                'metrics': {}
            }
        
        total_chars = len(text)
        reasons = []
        
        # Basic character analysis
        spaces = text.count(' ')
        alpha_chars = sum(1 for c in text if c.isalpha())
        digit_chars = sum(1 for c in text if c.isdigit())
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
        non_ascii_chars = sum(1 for c in text if ord(c) > 127)
        
        # Word analysis
        words = text.split()
        word_count = len(words)
        
        if word_count == 0:
            return {
                'is_likely_scanned': True,
                'quality_state': 'blocked', 
                'reasons': ['No words found'],
                'metrics': {}
            }
        
        # Calculate quality metrics
        avg_chars_per_page = total_chars / pages if pages > 0 else 0
        whitespace_ratio = spaces / total_chars if total_chars > 0 else 0
        special_ratio = special_chars / total_chars if total_chars > 0 else 0
        non_ascii_ratio = non_ascii_chars / total_chars if total_chars > 0 else 0
        
        # Word quality metrics
        single_char_words = sum(1 for word in words if len(word) == 1 and word.isalpha())
        short_words = sum(1 for word in words if len(word) <= 2)
        words_with_digits = sum(1 for word in words if any(c.isdigit() for c in word))
        
        prop_single_char_words = single_char_words / word_count
        prop_short_words = short_words / word_count
        prop_words_with_digits = words_with_digits / word_count
        avg_word_len = sum(len(word) for word in words) / word_count
        
        # OCR artifact detection
        hyphen_count = text.count('-')
        hyphen_rate = (hyphen_count / total_chars) * 1000  # per 1000 chars
        
        # Multi-space runs (common in OCR)
        multi_space_runs = len(re.findall(r'\s{2,}', text))
        multi_space_rate = (multi_space_runs / total_chars) * 1000
        
        # Broken word patterns (like "thi~ ~s" or "wo rd")
        broken_patterns = len(re.findall(r'[A-Za-z][\s~\-]{1,3}[A-Za-z]', text))
        broken_pattern_rate = (broken_patterns / total_chars) * 1000
        
        # RELAXED THRESHOLDS - Require multiple metrics to fail for rejection
        
        # Volume-based (relaxed for invoices/payment docs)
        if avg_chars_per_page < 200:
            reasons.append(f'Very low text density ({avg_chars_per_page:.0f} chars/page < 200)')
        
        # Quality-based thresholds (much more reasonable)
        if special_ratio > 0.35:
            reasons.append(f'High special character ratio ({special_ratio:.2%} > 35%)')
            
        if whitespace_ratio > 0.45:
            reasons.append(f'Excessive whitespace ({whitespace_ratio:.2%} > 45%)')
            
        if non_ascii_ratio > 0.08:
            reasons.append(f'Non-ASCII artifacts ({non_ascii_ratio:.2%} > 8%)')
            
        if prop_single_char_words > 0.20:
            reasons.append(f'Too many single-character words ({prop_single_char_words:.2%} > 20%)')
            
        if prop_short_words > 0.45:
            reasons.append(f'Too many short words ({prop_short_words:.2%} > 45%)')
            
        if avg_word_len < 3.0:
            reasons.append(f'Very short average word length ({avg_word_len:.1f} < 3.0)')
            
        if prop_words_with_digits > 0.35:
            reasons.append(f'Excessive digit-word mix ({prop_words_with_digits:.2%} > 35%)')
            
        if hyphen_rate > 80:
            reasons.append(f'Excessive hyphens ({hyphen_rate:.1f} per 1000 chars > 80)')
            
        if multi_space_rate > 25:
            reasons.append(f'Multiple space runs ({multi_space_rate:.1f} per 1000 chars > 25)')
            
        if broken_pattern_rate > 50:
            reasons.append(f'Broken word patterns ({broken_pattern_rate:.1f} per 1000 chars > 50)')
        
        # Severity override for extreme cases (immediate rejection)
        extreme_issues = []
        if broken_pattern_rate > 150:
            extreme_issues.append(f'Severely broken text patterns ({broken_pattern_rate:.1f} per 1000 chars)')
        if non_ascii_ratio > 0.25:
            extreme_issues.append(f'Extreme non-ASCII ratio ({non_ascii_ratio:.2%})')
        if multi_space_rate > 80:
            extreme_issues.append(f'Extreme spacing issues ({multi_space_rate:.1f} per 1000 chars)')
        
        # Reject immediately if extreme issues OR require 2+ normal issues
        is_scanned = len(extreme_issues) > 0 or len(reasons) >= 2
        
        # Add extreme issues to reasons for user feedback
        if extreme_issues:
            reasons = extreme_issues + reasons
        
        # Determine quality state (3-tier system)
        quality_state = self._determine_quality_state(extreme_issues, reasons, {
            'broken_pattern_rate': broken_pattern_rate,
            'non_ascii_ratio': non_ascii_ratio,
            'multi_space_rate': multi_space_rate,
            'avg_chars_per_page': avg_chars_per_page
        })
        
        return {
            'is_likely_scanned': is_scanned,
            'quality_state': quality_state,
            'reasons': reasons[:3],  # Show first 3 reasons to user
            'metrics': {
                'avg_chars_per_page': avg_chars_per_page,
                'whitespace_ratio': whitespace_ratio,
                'special_ratio': special_ratio,
                'non_ascii_ratio': non_ascii_ratio,
                'prop_single_char_words': prop_single_char_words,
                'prop_short_words': prop_short_words,
                'avg_word_len': avg_word_len,
                'prop_words_with_digits': prop_words_with_digits,
                'hyphen_rate': hyphen_rate,
                'multi_space_rate': multi_space_rate,
                'broken_pattern_rate': broken_pattern_rate
            }
        }
    
    def _determine_quality_state(self, extreme_issues: List[str], reasons: List[str], metrics: Dict[str, float]) -> str:
        """Determine document quality state: good, degraded, or blocked"""
        
        # Blocked: Extreme issues or 2+ normal issues
        if len(extreme_issues) > 0 or len(reasons) >= 2:
            return "blocked"
        
        # Degraded: Any single issue OR moderate quality concerns
        if len(reasons) > 0:
            return "degraded"
            
        # Check for moderate degraded thresholds (80% of extreme levels)
        degraded_thresholds = {
            'broken_pattern_rate': 120,  # 80% of 150
            'non_ascii_ratio': 0.20,     # 80% of 0.25
            'multi_space_rate': 64,      # 80% of 80
            'avg_chars_per_page': 250    # Moderate concern level
        }
        
        # Check if any metric is in degraded range
        for metric, threshold in degraded_thresholds.items():
            if metric in metrics:
                if metric == 'avg_chars_per_page':
                    if metrics[metric] < threshold:
                        return "degraded"
                else:
                    if metrics[metric] > threshold:
                        return "degraded"
        
        return "good"
    
    def _get_scanned_pdf_message(self, reasons = None, filename = None) -> str:
        """Return user-friendly message for scanned PDF detection with specific reasons"""
        filename_text = f': "{filename}"' if filename else ''
        base_message = (
            f"🔍 **Scanned/Image-Based PDF Detected{filename_text}**\n\n"
            "This PDF appears to be scanned or image-based and cannot be processed directly.\n\n"
        )
        
        # Add specific reasons if available
        if reasons:
            reasons_text = "**Issues Detected:**\n"
            for reason in reasons:
                reasons_text += f"• {reason}\n"
            reasons_text += "\n"
        else:
            reasons_text = ""
        
        solutions = (
            "**Quick Solutions:**\n\n"
            "- **ChatGPT-4 Vision**: Upload your PDF to ChatGPT-4 with Vision and ask it to convert to text, then paste into a new document\n\n"
            "- **OCR Software**: Use Adobe Acrobat, Google Docs, or other OCR tools to convert to searchable text\n\n"
            "- **Contact Support**: We can assist with document processing guidance\n\n"
            "**Once converted to text-based PDF, please re-upload for analysis.**"
        )
        
        return base_message + reasons_text + solutions