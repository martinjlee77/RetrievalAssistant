"""
LLM Utilities - Following Streamlit Best Practices
Handles OpenAI API calls, error handling, caching, and file conversion utilities
"""
import os
import streamlit as st
from openai import OpenAI, AsyncOpenAI
from typing import Dict, List, Any, Optional, Union, cast
import json
import time
import io
from docx import Document
from fpdf import FPDF

# Initialize OpenAI client
@st.cache_resource
def get_openai_client():
    """Initialize and cache OpenAI client"""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        st.error("OpenAI API key not found. Please check your secrets configuration.")
        st.stop()
    return OpenAI(api_key=api_key)

async def make_llm_call_async(
    client,
    messages: List[Dict[str, str]],  # CHANGED from prompt: str
    temperature: float = 0.3,
    max_tokens: Optional[int] = None,
    model: str = "gpt-5",  # User requested gpt-5 as the default model
    response_format: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    Asynchronous version of make_llm_call for concurrent execution
    """
    try:
        # Create async client for this call
        api_key = os.environ.get("OPENAI_API_KEY")
        async_client = AsyncOpenAI(api_key=api_key)
        
        # Prepare request parameters
        request_params = {
            "model": model,
            "messages": messages,  # This now uses the passed-in messages list
            "temperature": temperature
        }
        
        if max_tokens:
            # Use correct parameter name based on model
            if model.startswith("gpt-5") or model.startswith("o1"):
                request_params["max_completion_tokens"] = max_tokens
            else:
                request_params["max_tokens"] = max_tokens
        if response_format:
            request_params["response_format"] = response_format
        
        # Make async API call
        response = await async_client.chat.completions.create(**request_params)
        return response.choices[0].message.content
        
    except Exception as e:
        # Log error internally without displaying technical details to users
        import logging
        logging.error(f"Async LLM API call failed: {str(e)}")
        return None

def make_llm_call(
    client,
    messages: List[Dict[str, str]],  # CHANGED from prompt: str
    temperature: float = 0.3,
    max_tokens: Optional[int] = None,
    model: str = "gpt-5",  # User requested gpt-5 as the default model
    response_format: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    Make LLM API call with error handling and rate limiting
    Following Streamlit best practices for API management
    """
    try:
        with st.spinner("Analyzing with AI..."):
            # Messages are now passed in directly
            
            # Prepare request parameters
            request_params = {
                "model": model,
                "messages": messages,
                "temperature": temperature,
            }
            
            # Add optional parameters if provided
            if response_format:
                request_params["response_format"] = response_format
            if max_tokens:
                # Use correct parameter name based on model
                if model.startswith("gpt-5") or model.startswith("o1"):
                    request_params["max_completion_tokens"] = max_tokens
                else:
                    request_params["max_tokens"] = max_tokens
                
            response = client.chat.completions.create(**request_params)
        return response.choices[0].message.content
    
    except Exception as e:
        handle_llm_error(e)
        return None

def handle_llm_error(error: Exception):
    """Handle LLM API errors with internal logging only - no user-facing technical details"""
    import logging
    error_message = str(error).lower()
    
    # Log all errors internally for debugging
    logging.error(f"LLM API error: {str(error)}")
    
    # Only display user-facing errors for configuration issues, not temporary API problems
    if "quota" in error_message or "billing" in error_message:
        st.error("💳 API quota exceeded. Please check your OpenAI billing settings.")
    elif "invalid api key" in error_message or "unauthorized" in error_message:
        st.error("🔑 Invalid API key. Please check your OpenAI API key configuration.")
    elif "context length" in error_message or "token" in error_message:
        st.error("📄 Content too long. Please try with a shorter document or input.")
    # Rate limits, server errors, and other temporary issues are handled silently
    # The system has retry logic and graceful degradation

@st.cache_data(ttl=3600)  # Cache for 1 hour
def cached_llm_call(
    prompt: str, 
    system_message: str = None,
    model: str = "gpt-5",
    temperature: float = 0.3
) -> Optional[str]:
    """
    Cached LLM call for frequently requested analysis
    Follows Streamlit best practices for caching expensive operations
    """
    messages = []
    if system_message:
        messages.append({"role": "system", "content": system_message})
    messages.append({"role": "user", "content": prompt})
    
    result = make_llm_call(messages, model, temperature)
    return result if result else ""

def stream_llm_response(
    messages: List[Dict[str, str]], 
    model: str = "gpt-5",
    temperature: float = 0.3
):
    """
    Stream LLM response for real-time display
    Uses st.write_stream for token-by-token display
    """
    client = get_openai_client()
    
    try:
        # Cast messages to proper type for OpenAI
        openai_messages = cast(List[Any], messages)
        # Prepare request parameters with model-specific parameter handling
        request_params = {
            "model": model,
            "messages": openai_messages,
            "temperature": temperature,
            "stream": True
        }
        
        response = client.chat.completions.create(**request_params)
        
        def response_generator():
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    yield chunk.choices[0].delta.content
        
        return st.write_stream(response_generator())
    
    except Exception as e:
        handle_llm_error(e)
        return None

def get_model_options() -> Dict[str, str]:
    """Get available model options for debugging UI"""
    return {
        "GPT-5 (Latest)": "gpt-5",
        "GPT-5 Mini": "gpt-5-mini",
        "GPT-4 Turbo": "gpt-4-turbo-preview"
    }

def create_debug_sidebar():
    """
    Create debugging sidebar for prompt engineering
    Following Streamlit best practices for experimentation
    """
    with st.sidebar:
        st.subheader("🔧 AI Debug Controls")
        
        model_options = get_model_options()
        selected_model = st.selectbox(
            "Model",
            options=list(model_options.keys()),
            index=0
        )
        
        temperature = st.slider(
            "Temperature",
            min_value=0.0,
            max_value=2.0,
            value=0.3,
            step=0.1,
            help="Higher values make output more creative but less focused"
        )
        
        max_tokens = st.slider(
            "Max Tokens",
            min_value=1000,
            max_value=8192,
            value=4096,
            step=256,
            help="Maximum response length"
        )
        
        return {
            "model": model_options[selected_model],
            "temperature": temperature,
            "max_tokens": max_tokens
        }

def extract_contract_terms(client, contract_text: str, step_context: str = "comprehensive_analysis") -> List[str]:
    """
    Extract key contract terms for enhanced RAG search
    
    Args:
        client: OpenAI client
        contract_text: Full contract text
        step_context: Analysis step context
        
    Returns:
        List of extracted contract terms for semantic search
    """
    try:
        prompt = f"""
        Extract 5-7 key terms and phrases from this contract that would be most relevant for ASC 606 revenue recognition analysis:
        
        CONTRACT TEXT:
        {contract_text[:3000]}...
        
        Focus on:
        - Service/product descriptions (e.g., SaaS license, implementation)
        - Performance obligations
        - Payment terms (e.g., Net 30, upfront payment)
        - Contract timing/duration
        - Acceptance clauses, warranties, or return rights
        - Discounts, rebates, penalties, or other variable consideration
        - Customer responsibilities
        
        Return only the key terms, one per line, no explanations.
        """
        
        # NEW: Wrap prompt in messages list structure
        messages = [{"role": "user", "content": prompt}]
        response = make_llm_call(
            client=client,
            messages=messages,  # Pass the new messages list
            model="gpt-5-mini",
            max_tokens=200,
            temperature=0.3
        )
        
        if response:
            # Split response into individual terms
            terms = [term.strip() for term in response.split('\n') if term.strip()]
            return terms[:7]  # Limit to 7 terms
        
        return []
        
    except Exception as e:
        st.warning(f"Contract term extraction failed: {e}")
        return []

def validate_api_key() -> bool:
    """Validate OpenAI API key is properly configured"""
    try:
        client = get_openai_client()
        # Test with minimal API call
        openai_messages = cast(List[Any], [{"role": "user", "content": "test"}])
        client.chat.completions.create(
            model="gpt-5-mini",
            messages=openai_messages,
            max_tokens=5
        )
        return True
    except Exception:
        return False

def create_docx_from_text(text_content, contract_data=None):
    """Creates a professional accounting memo in DOCX format with rule-based parsing and Segoe UI font."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.parser import OxmlElement
    from docx.oxml.ns import qn
    from datetime import datetime
    import re
    import logging
    import hashlib
    from functools import lru_cache
    
    document = Document()
    
    # === PHASE 1: PROFESSIONAL DOCUMENT STRUCTURE ===
    
    # Set default font to match HTML memo styling
    style = document.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'  # Matches HTML system font
    font.size = Pt(12)  # Matches HTML 12pt
    font.color.rgb = RGBColor(0, 0, 0)  # Black text like HTML
    
    # Configure paragraph spacing to match HTML (12pt consistent spacing)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(12)  # Matches HTML consistent_spacing
    paragraph_format.line_spacing = 1.0  # Single spacing as requested
    
    # Set standard accounting memo margins (1" all sides)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header with Controller.cpa branding
    header = document.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "VeritasLogic.ai"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.runs[0]
    header_run.font.name = 'Segoe UI'
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(0, 0, 0)  # Black only
    
    # Add footer with page numbers
    footer = document.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
    footer_run.font.name = 'Segoe UI'
    footer_run.font.size = Pt(10)
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    footer_run._r.append(fldChar1)
    footer_run._r.append(instrText)
    footer_run._r.append(fldChar2)
    

    
    # === PHASE 2: CONTENT PARSING AND FORMATTING (ENHANCED) ===
    # This enhanced parser understands Markdown headings, bold, tables, and lists,
    # and applies the document's pre-defined styles to match HTML output quality.
    
    # === PHASE 2: RULE-BASED PARSING SYSTEM ===
    # Implement robust, maintainable parsing using regex rules
    
    def configure_heading_styles():
        """Configure custom heading styles with enhanced professional formatting"""
        try:
            # Heading 1 - Main sections (matches HTML header styling)
            heading1_style = document.styles.add_style('Custom Heading 1', 1)
            heading1_style.font.name = 'Segoe UI'
            heading1_style.font.size = Pt(16)  # Matches HTML header_size
            heading1_style.font.bold = True
            heading1_style.font.color.rgb = RGBColor(0, 0, 0)  # Black like HTML
            heading1_style.paragraph_format.space_before = Pt(16)  # More consistent spacing
            heading1_style.paragraph_format.space_after = Pt(12)  # Matches HTML consistent_spacing
            heading1_style.paragraph_format.keep_with_next = True
        except Exception as e:
            logging.warning(f"Could not create Custom Heading 1 style: {e}")
            
        try:
            # Heading 2 - Subsections (matches HTML subheader styling)
            heading2_style = document.styles.add_style('Custom Heading 2', 1)
            heading2_style.font.name = 'Segoe UI'
            heading2_style.font.size = Pt(14)  # Matches HTML subheader_size
            heading2_style.font.bold = True
            heading2_style.font.color.rgb = RGBColor(0, 0, 0)  # Black like HTML
            heading2_style.paragraph_format.space_before = Pt(12)  # Consistent with HTML
            heading2_style.paragraph_format.space_after = Pt(12)  # Matches HTML consistent_spacing
            heading2_style.paragraph_format.keep_with_next = True
        except Exception as e:
            logging.warning(f"Could not create Custom Heading 2 style: {e}")
            
        try:
            # Heading 3 - Sub-subsections  
            heading3_style = document.styles.add_style('Custom Heading 3', 1)
            heading3_style.font.name = 'Segoe UI'
            heading3_style.font.size = Pt(12)
            heading3_style.font.bold = True
            heading3_style.font.color.rgb = RGBColor(0, 0, 0)  # Black only
            heading3_style.paragraph_format.space_before = Pt(8)
            heading3_style.paragraph_format.space_after = Pt(4)
        except Exception as e:
            logging.warning(f"Could not create Custom Heading 3 style: {e}")
            
        try:
            # NEW: Subsection Header Style for OVERALL CONCLUSION, KEY FINDINGS, etc.
            subsection_style = document.styles.add_style('Subsection Header', 1)
            subsection_style.font.name = 'Segoe UI'
            subsection_style.font.size = Pt(11)
            subsection_style.font.bold = True
            subsection_style.font.color.rgb = RGBColor(0, 0, 0)
            subsection_style.font.all_caps = True  # UPPERCASE styling
            subsection_style.paragraph_format.space_before = Pt(12)
            subsection_style.paragraph_format.space_after = Pt(6)
        except Exception as e:
            logging.warning(f"Could not create Subsection Header style: {e}")
    
    # Compile regex patterns once for performance - ENHANCED FOR NESTED STRUCTURES
    HEADING1_PATTERN = re.compile(r'^#\s+(.*?)$')
    HEADING2_PATTERN = re.compile(r'^##\s+(.*?)$') 
    HEADING3_PATTERN = re.compile(r'^###\s+(.*?)$')
    SUBSECTION_PATTERN = re.compile(r'^(OVERALL CONCLUSION|KEY FINDINGS|CONTRACT DATA SUMMARY|DOCUMENTS REVIEWED|DETAILED ANALYSIS|CONCLUSION)$')
    
    # Enhanced bullet patterns to capture indentation and different markers
    NESTED_BULLET_PATTERN = re.compile(r'^(\s+)[\*\-•]\s+(.*?)$')  # Captures leading whitespace + bullet char
    SUB_BULLET_PATTERN = re.compile(r'^-\s+(.*?)$')               # Hyphen as sub-bullet marker  
    MAIN_BULLET_PATTERN = re.compile(r'^[\*•]\s+(.*?)$')          # Asterisk/bullet as main bullet
    BULLET_PATTERN = re.compile(r'^[\*\-•]\s+(.*?)$')            # Fallback for any bullet
    
    NUMBERED_PATTERN = re.compile(r'^\d+\.\s+(.*?)$')
    BLOCKQUOTE_PATTERN = re.compile(r'^>\s*(.*?)$')
    HORIZONTAL_RULE_PATTERN = re.compile(r'^---+$')
    TABLE_PLACEHOLDER_PATTERN = re.compile(r'^\[TABLE_PLACEHOLDER\]$')
    
    def add_heading1(doc, match):
        """Add heading 1 with custom style"""
        try:
            doc.add_paragraph(match.group(1), style='Custom Heading 1')
        except Exception:
            doc.add_paragraph(match.group(1), style='Heading 1')
    
    def add_heading2(doc, match):
        """Add heading 2 with custom style and proper spacing"""
        # A6: Add more space before main section headers
        heading_text = match.group(1)
        if any(x in heading_text for x in ["EXECUTIVE SUMMARY", "CONTRACT OVERVIEW", "DETAILED ASC 606", "KEY PROFESSIONAL", "FINANCIAL IMPACT", "CONCLUSION"]):
            doc.add_paragraph()  # Extra space before main sections
        
        try:
            doc.add_paragraph(heading_text, style='Custom Heading 2')
        except Exception:
            doc.add_paragraph(heading_text, style='Heading 2')
    
    def add_heading3(doc, match):
        """Add heading 3 with custom style and reduced spacing"""
        heading_text = match.group(1)
        # A7: Reduce space before "Detailed Analysis" and similar subsections
        # A8: Remove redundant "Conclusion:" from headers
        if heading_text.startswith("Conclusion:"):
            heading_text = heading_text.replace("Conclusion:", "").strip()
        
        try:
            doc.add_paragraph(heading_text, style='Custom Heading 3')
        except Exception:
            doc.add_paragraph(heading_text, style='Heading 3')
    
    def add_bullet_point(doc, match):
        """Add bullet point with intelligent structure-based indentation"""
        bullet_text = match.group(1)
        bullet_para = doc.add_paragraph()
        _parse_text_formatting(bullet_para, bullet_text)
        bullet_para.style = 'List Bullet'
        
        # ROBUST APPROACH: Detect sub-bullet patterns structurally
        # Method 1: Check for leading spaces/tabs in the original line (nested markdown)
        full_line = match.group(0)  # Get the entire matched line including bullet marker
        if full_line.startswith('  ') or full_line.startswith('\t'):
            # This is a nested bullet with leading whitespace
            bullet_para.paragraph_format.left_indent = Inches(0.5)
            return
        
        # Method 2: Detect sub-bullet markers (e.g., "- " for sub-bullets vs "• " for main)
        if full_line.strip().startswith('- '):
            # Using hyphen as sub-bullet indicator
            bullet_para.paragraph_format.left_indent = Inches(0.5)
            return
            
        # Method 3: Contextual analysis - check if this follows a "parent" concept
        # Look for patterns that indicate this is a breakdown of the previous item
        sub_bullet_indicators = [
            ':',  # Previous line ended with colon (introducing a list)
            'obligations:',  # Performance obligations breakdown
            'components:',   # Component breakdown
            'includes:',     # Inclusion list
            'consists of:',  # Composition breakdown
        ]
        
        # Get the last few paragraphs to analyze context
        recent_paragraphs = doc.paragraphs[-3:] if len(doc.paragraphs) >= 3 else doc.paragraphs
        context_text = ' '.join([p.text.lower() for p in recent_paragraphs])
        
        if any(indicator in context_text for indicator in sub_bullet_indicators):
            bullet_para.paragraph_format.left_indent = Inches(0.5)
    
    def add_numbered_item(doc, match):
        """Add numbered list item with formatting support"""
        num_para = doc.add_paragraph()
        _parse_text_formatting(num_para, match.group(1))
        num_para.style = 'List Number'
    
    def add_subsection_header(doc, match):
        """Add professional subsection header (OVERALL CONCLUSION, KEY FINDINGS, etc.)"""
        try:
            doc.add_paragraph(match.group(1), style='Subsection Header')
        except Exception:
            # Fallback to manual formatting
            para = doc.add_paragraph()
            run = para.add_run(match.group(1))
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black only
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
    
    def add_blockquote(doc, match):
        """Add enhanced blockquote with professional contract citation styling"""
        quote_para = doc.add_paragraph()
        _parse_text_formatting(quote_para, match.group(1))
        # Enhanced professional styling for contract excerpts
        quote_para.paragraph_format.left_indent = Inches(0.5)
        quote_para.paragraph_format.right_indent = Inches(0.5)
        quote_para.paragraph_format.space_before = Pt(6)
        quote_para.paragraph_format.space_after = Pt(6)
        # Add subtle border effect through shading
        try:
            from docx.oxml.shared import qn
            from docx.oxml import parse_xml
            shading_elm = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(qn('w:val')))
            quote_para._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            pass  # Graceful fallback if shading fails
        
        for run in quote_para.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(85, 85, 85)
    
    def add_horizontal_rule(doc, match):
        """Add section separator"""
        doc.add_paragraph()
    
    def add_table_placeholder(doc, match):
        """Handle table placeholder - will be processed by main loop"""
        return "TABLE_PLACEHOLDER"  # Signal to main loop
    
    configure_heading_styles()
    
    # Preprocess content to handle HTML-like formatting and tables
    def preprocess_content(content):
        """Preprocess content to extract tables and clean formatting"""
        # The header is now passed in and does not need to be removed.
        
        # Extract and store contract overview tables (with bold headers)
        table_pattern = r'\|\s*\*\*.*?\*\*\s*\|.*?\|\s*\n(?:\|.*?\|\s*\n)*'
        tables = re.findall(table_pattern, content, re.MULTILINE | re.DOTALL)
        
        # Remove tables from content for now, we'll add them back specially
        for table in tables:
            content = content.replace(table, '\n[TABLE_PLACEHOLDER]\n')
        
        return content, tables
    
    processed_content, extracted_tables = preprocess_content(text_content)
    
    # Split the entire text_content into individual lines for processing
    lines = processed_content.split('\n')
    table_index = 0
    
    def add_nested_bullet_point(doc, match):
        """Handle nested bullets with preserved indentation"""
        whitespace = match.group(1)  # Captured leading whitespace
        bullet_text = match.group(2)
        bullet_para = doc.add_paragraph()
        _parse_text_formatting(bullet_para, bullet_text)
        bullet_para.style = 'List Bullet'
        
        # Calculate indentation based on whitespace amount
        indent_level = len(whitespace.expandtabs(4)) // 4  # Convert tabs to spaces, then count levels
        base_indent = 0.5 * indent_level  # 0.5 inches per level
        bullet_para.paragraph_format.left_indent = Inches(base_indent)
    
    def add_sub_bullet_point(doc, match):
        """Handle sub-bullets marked with hyphens"""
        bullet_text = match.group(1)
        bullet_para = doc.add_paragraph()
        _parse_text_formatting(bullet_para, bullet_text)
        bullet_para.style = 'List Bullet'
        bullet_para.paragraph_format.left_indent = Inches(0.5)  # Standard sub-bullet indent
    
    def add_main_bullet_point(doc, match):
        """Handle main bullets (asterisk/bullet points)"""
        bullet_text = match.group(1)
        bullet_para = doc.add_paragraph()
        _parse_text_formatting(bullet_para, bullet_text)
        bullet_para.style = 'List Bullet'
        # No additional indentation for main bullets

    # Define parsing rules in order of specificity (MOST SPECIFIC FIRST)
    parse_rules = [
        (HEADING1_PATTERN, add_heading1),
        (HEADING2_PATTERN, add_heading2), 
        (HEADING3_PATTERN, add_heading3),
        (SUBSECTION_PATTERN, add_subsection_header),
        (NESTED_BULLET_PATTERN, add_nested_bullet_point),  # Check nested bullets first
        (SUB_BULLET_PATTERN, add_sub_bullet_point),        # Then sub-bullet markers
        (MAIN_BULLET_PATTERN, add_main_bullet_point),      # Then main bullets
        (BULLET_PATTERN, add_bullet_point),                # Fallback for any remaining bullets
        (NUMBERED_PATTERN, add_numbered_item),
        (BLOCKQUOTE_PATTERN, add_blockquote),
        (HORIZONTAL_RULE_PATTERN, add_horizontal_rule),
        (TABLE_PLACEHOLDER_PATTERN, add_table_placeholder)
    ]
    
    @lru_cache(maxsize=256)
    def _parse_text_formatting_cached(text_hash: str, text: str):
        """Cached helper for text formatting parsing"""
        return _parse_text_formatting_logic(text)
        
    def _parse_text_formatting_logic(text: str):
        """Core text formatting logic - separated for caching"""
        if not text:
            return []
            
        # Handle bold (**text**) and italic (*text*) formatting
        parts = []
        current_pos = 0
        
        # Find all bold and italic markers
        bold_pattern = re.compile(r'\*\*([^*]+?)\*\*')
        italic_pattern = re.compile(r'\*([^*]+?)\*')
        
        # Process bold formatting first
        for match in bold_pattern.finditer(text):
            # Add text before the match
            if match.start() > current_pos:
                parts.append(('normal', text[current_pos:match.start()]))
            # Add bold text
            parts.append(('bold', match.group(1)))
            current_pos = match.end()
        
        # Add remaining text
        if current_pos < len(text):
            remaining_text = text[current_pos:]
            # Check for italic in remaining text
            italic_pos = 0
            for match in italic_pattern.finditer(remaining_text):
                if match.start() > italic_pos:
                    parts.append(('normal', remaining_text[italic_pos:match.start()]))
                parts.append(('italic', match.group(1)))
                italic_pos = match.end()
            if italic_pos < len(remaining_text):
                parts.append(('normal', remaining_text[italic_pos:]))
        
        # If no formatting found, treat as normal text
        if not parts:
            parts.append(('normal', text))
        
        return parts
    
    def _parse_text_formatting(paragraph, text):
        """Enhanced text formatting parser supporting bold and italic with caching"""
        if not text:
            return
        
        # Use cached formatting parsing
        text_hash = hashlib.md5(text.encode()).hexdigest()
        parts = _parse_text_formatting_cached(text_hash, text)
        
        # Apply formatting to paragraph
        for format_type, content in parts:
            if not content:
                continue
            run = paragraph.add_run(content)
            run.font.name = 'Segoe UI'  # Ensure consistent font
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
    
    def _add_contract_table(doc, table_markdown):
        """Enhanced table parser with special handling for journal entries in DOCX"""
        if not table_markdown or not table_markdown.strip():
            return
            
        lines = [line.strip() for line in table_markdown.strip().split('\n') if line.strip()]
        if len(lines) < 2:
            return
        
        # Parse headers - more robust parsing
        header_line = lines[0]
        if not '|' in header_line:
            return
            
        headers = []
        for cell in header_line.split('|'):
            cleaned_cell = cell.strip().strip('*').strip()
            if cleaned_cell:
                headers.append(cleaned_cell)
        
        if not headers:
            return
        
        # Note: Journal entries now use simple text format, not tables
        
        # Find separator line (|---|---|) and data rows
        data_rows = []
        separator_found = False
        
        for line in lines[1:]:
            if not separator_found and re.match(r'^[\|\-\s:]+$', line):
                separator_found = True
                continue
            
            if '|' in line and separator_found:
                row_data = []
                for cell in line.split('|'):
                    cleaned_cell = cell.strip().strip('*').strip()
                    if cleaned_cell:
                        row_data.append(cleaned_cell)
                
                # Only add row if it has the right number of columns
                if len(row_data) >= len(headers):
                    data_rows.append(row_data[:len(headers)])
        
        if not data_rows:
            return
        
        # Add missing fields for contract tables
        if len(headers) == 2 and any("Element" in h or "Details" in h for h in headers):
            existing_rows = [row_data[0] for row_data in data_rows if row_data]
            if "Document Classification" not in str(existing_rows):
                data_rows.extend([
                    ["Document Classification", "Internal Use Only"],
                    ["Review Status", "Preliminary Analysis"]
                ])
        
        # Use smart table creation with proper formatting
        try:
            from utils.table_helpers import create_smart_table
            create_smart_table(doc, headers, data_rows)
        except ImportError:
            # Fallback to basic table creation if import fails
            table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Basic formatting
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                if i < len(header_row.cells):
                    header_row.cells[i].text = header
            
            for row_idx, row_data in enumerate(data_rows):
                table_row = table.rows[row_idx + 1]
                for col_idx, cell_data in enumerate(row_data):
                    if col_idx < len(table_row.cells):
                        table_row.cells[col_idx].text = str(cell_data)
        
        doc.add_paragraph()  # Add spacing after table
    
    def _add_journal_entry_list(doc, table_markdown):
        """Convert journal entry table to clean, professional list format for DOCX"""
        lines = [line.strip() for line in table_markdown.strip().split('\n') if line.strip()]
        if len(lines) < 3:  # Need header, separator, and at least one data row
            return
        
        # Parse the table structure
        header_line = lines[0]
        headers = [cell.strip().strip('*').strip() for cell in header_line.split('|') if cell.strip().strip('*').strip()]
        
        # Find data rows (skip separator line)
        data_rows = []
        separator_found = False
        
        for line in lines[1:]:
            if not separator_found and re.match(r'^[\|\-\s:]+$', line):
                separator_found = True
                continue
            
            if '|' in line and separator_found:
                row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                if len(row_data) >= len(headers):
                    data_rows.append(row_data[:len(headers)])
        
        if not data_rows:
            return
        
        # Add section header for journal entries
        header_para = doc.add_paragraph()
        header_para.paragraph_format.space_before = Pt(6)
        header_para.paragraph_format.space_after = Pt(12)
        header_run = header_para.add_run("Illustrative Journal Entries")
        header_run.font.name = 'Segoe UI'
        header_run.font.bold = True
        header_run.font.size = Pt(12)
        header_run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black color
        
        # Group entries by transaction date
        current_date = None
        transaction_total_debit = 0
        transaction_total_credit = 0
        
        for row in data_rows:
            if len(row) < 3:  # Need at least Date, Account, and one amount
                continue
                
            date_val = row[0] if row[0] and '[date]' not in row[0].lower() else "Contract Date"
            account = row[1] if len(row) > 1 else ""
            debit = row[2] if len(row) > 2 else ""
            credit = row[3] if len(row) > 3 else ""
            description = row[4] if len(row) > 4 else ""
            
            # Skip empty rows or description-only rows
            if not account or account.lower().startswith('to record'):
                continue
            
            # Start new transaction entry
            if date_val != current_date:
                # Close previous transaction with totals if needed
                if current_date and (transaction_total_debit > 0 or transaction_total_credit > 0):
                    doc.add_paragraph()  # Spacing between transactions
                
                current_date = date_val
                transaction_total_debit = 0
                transaction_total_credit = 0
                
                # Transaction date header
                date_para = doc.add_paragraph()
                date_para.paragraph_format.space_before = Pt(12)
                date_para.paragraph_format.space_after = Pt(6)
                date_run = date_para.add_run(f"{date_val}:")
                date_run.font.name = 'Segoe UI'
                date_run.font.bold = True
                date_run.font.size = Pt(11)
            
            # Clean up amounts
            debit_clean = debit.strip().replace('$', '').replace(',', '') if debit and debit.strip() and debit.strip() != '-' else ''
            credit_clean = credit.strip().replace('$', '').replace(',', '') if credit and credit.strip() and credit.strip() != '-' else ''
            
            # Format journal entry line
            entry_para = doc.add_paragraph()
            entry_para.paragraph_format.left_indent = Inches(0.3)
            entry_para.paragraph_format.space_after = Pt(3)
            
            # Create clean, readable format
            if debit_clean:
                # Debit entry
                amount_formatted = f"${debit_clean}" if not debit.startswith('$') else debit
                entry_line = f"{account:<40} {amount_formatted:>12}"
                # Track totals
                try:
                    transaction_total_debit += float(debit_clean)
                except (ValueError, TypeError):
                    pass
            elif credit_clean:
                # Credit entry (indented)
                amount_formatted = f"${credit_clean}" if not credit.startswith('$') else credit
                entry_line = f"    {account:<36} {amount_formatted:>12}"
                # Track totals
                try:
                    transaction_total_credit += float(credit_clean)
                except (ValueError, TypeError):
                    pass
            else:
                # No amount, just account name
                entry_line = f"{account}"
            
            # Add the formatted line
            if entry_line.strip():
                entry_run = entry_para.add_run(entry_line)
                entry_run.font.name = 'Consolas'  # Clean monospace font
                entry_run.font.size = Pt(10)
                entry_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add final spacing
        doc.add_paragraph()
        
        # Add note about entries format
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("Note: Journal entries exclude applicable sales tax and are presented for illustrative purposes.")
        note_run.font.name = 'Segoe UI'
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(64, 64, 64)

    # === ENHANCED RULE-BASED PARSING LOOP ===
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue  # Skip empty lines to avoid extra paragraphs
        
        # SPECIAL HANDLING: Document title formatting (support both ASC 606 and ASC 340-40)
        if ("TECHNICAL ACCOUNTING MEMORANDUM" in stripped_line or "ACCOUNTING POLICY MEMORANDUM" in stripped_line) and not stripped_line.startswith('#'):
            title_para = document.add_paragraph(stripped_line)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue
        
        # Handle table placeholders specially
        if stripped_line == '[TABLE_PLACEHOLDER]' and table_index < len(extracted_tables):
            _add_contract_table(document, extracted_tables[table_index])
            table_index += 1
            continue
        
        # Apply parsing rules in order of specificity
        was_handled = False
        for pattern, handler in parse_rules:
            match = pattern.match(stripped_line)
            if match:
                result = handler(document, match)
                # Special handling for table placeholders
                if result == "TABLE_PLACEHOLDER" and table_index < len(extracted_tables):
                    _add_contract_table(document, extracted_tables[table_index])
                    table_index += 1
                was_handled = True
                break
        
        # If no rule matched, treat as normal paragraph
        if not was_handled:
            p = document.add_paragraph()
            _parse_text_formatting(p, stripped_line)
    
    # === PHASE 3: SIMPLIFIED METADATA SECTION ===
    
    # Add spacing and simplified metadata (no page break)
    document.add_paragraph()
    document.add_paragraph()
    
    # Simple metadata paragraphs instead of complex table
    current_date = datetime.now().strftime("%B %d, %Y")
    
    metadata_para = document.add_paragraph()
    metadata_run = metadata_para.add_run(f"Analysis Date: {current_date} | Review Status: Preliminary Analysis | Internal Use")
    metadata_run.font.name = 'Segoe UI'
    metadata_run.font.size = Pt(9)
    metadata_run.font.color.rgb = RGBColor(0, 0, 0)  # Black only
    metadata_para.alignment = 1  # Center alignment
    
    bio = io.BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

# Legacy PDF generation removed - replaced with WeasyPrint HTML-to-PDF approach
# See utils/html_export.py create_pdf_from_html() for new implementation
