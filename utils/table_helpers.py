"""
Professional table formatting helpers for DOCX generation.
Provides consistent, reusable table styling functions.
"""

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import parse_xml
import logging

def apply_professional_header_styling(cell):
    """Apply consistent header styling to table cells"""
    cell.text = cell.text.strip()
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = 'Lato'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Add blue header background
    try:
        shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(qn('w:val')))
        cell._element.get_or_add_tcPr().append(shading_elm)
    except Exception:
        pass  # Graceful fallback

def apply_data_cell_styling(cell, row_idx):
    """Apply consistent data cell styling with alternating row colors"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Lato'
            run.font.size = Pt(10)
    
    # Alternating row colors for better readability
    if row_idx % 2 == 0:
        try:
            shading_elm = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(qn('w:val')))
            cell._element.get_or_add_tcPr().append(shading_elm)
        except Exception:
            pass

def create_contract_table(doc, headers, data_rows):
    """Create and format contract data summary tables (2-column)"""
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Contract-specific column widths
    table.columns[0].width = Inches(2.2)  # Element column
    table.columns[1].width = Inches(4.3)  # Details column
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        if i < len(header_row.cells):
            header_row.cells[i].text = header
            apply_professional_header_styling(header_row.cells[i])
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table_row.cells):
                table_row.cells[col_idx].text = str(cell_data)
                apply_data_cell_styling(table_row.cells[col_idx], row_idx)
    
    return table

def create_journal_entry_table(doc, headers, data_rows):
    """Create and format journal entry tables (Account, Debit, Credit format)"""
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Journal entry specific column widths
    if len(headers) == 3:  # Account, Debit, Credit
        table.columns[0].width = Inches(3.5)  # Account (wider for descriptions)
        table.columns[1].width = Inches(1.5)  # Debit
        table.columns[2].width = Inches(1.5)  # Credit
    elif len(headers) == 4:  # Account, Description, Debit, Credit
        table.columns[0].width = Inches(2.0)  # Account
        table.columns[1].width = Inches(2.5)  # Description
        table.columns[2].width = Inches(1.25) # Debit
        table.columns[3].width = Inches(1.25) # Credit
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        if i < len(header_row.cells):
            header_row.cells[i].text = header
            apply_professional_header_styling(header_row.cells[i])
    
    # Add data rows with right-alignment for monetary columns
    for row_idx, row_data in enumerate(data_rows):
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table_row.cells):
                cell = table_row.cells[col_idx]
                cell.text = str(cell_data)
                apply_data_cell_styling(cell, row_idx)
                
                # Right-align monetary columns (last 2 columns)
                if col_idx >= len(headers) - 2:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    return table

def create_general_table(doc, headers, data_rows):
    """Create and format general tables with smart column sizing"""
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Smart column sizing based on content
    if len(headers) <= 3:
        # Equal distribution for small tables
        col_width = Inches(6.5) / len(headers)
        for col in table.columns:
            col.width = col_width
    else:
        # For larger tables, use proportional sizing
        table.autofit = False
        base_width = Inches(6.5) / len(headers)
        for col in table.columns:
            col.width = base_width
    
    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        if i < len(header_row.cells):
            header_row.cells[i].text = header
            apply_professional_header_styling(header_row.cells[i])
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        table_row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(table_row.cells):
                table_row.cells[col_idx].text = str(cell_data)
                apply_data_cell_styling(table_row.cells[col_idx], row_idx)
    
    return table

def detect_table_type(headers, data_rows):
    """Intelligently detect table type based on headers and content"""
    header_text = ' '.join(headers).lower()
    
    # Contract data table
    if any(term in header_text for term in ['element', 'details', 'contract', 'customer']):
        return 'contract'
    
    # Journal entry table
    if any(term in header_text for term in ['account', 'debit', 'credit', 'dr', 'cr']):
        return 'journal'
    
    # Default to general table
    return 'general'

def create_smart_table(doc, headers, data_rows):
    """Smart table creation that automatically applies appropriate formatting"""
    table_type = detect_table_type(headers, data_rows)
    
    if table_type == 'contract':
        return create_contract_table(doc, headers, data_rows)
    elif table_type == 'journal':
        return create_journal_entry_table(doc, headers, data_rows)
    else:
        return create_general_table(doc, headers, data_rows)