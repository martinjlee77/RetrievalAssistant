#!/usr/bin/env python3
"""
ASC 842 Leases Knowledge Base Seeding Script
Populates ChromaDB with ASC 842 authoritative and interpretative sources
Following proven ASC 340-40 pattern
"""

import os
import logging
import re
from pathlib import Path
from typing import List, Dict, Any
import chromadb
from chromadb.config import Settings
import chromadb.utils.embedding_functions as embedding_functions
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def clean_bullet_formatting(text: str) -> str:
    """
    Fix bullet formatting issues like 'aTopic' -> 'a. Topic' and 'bCosts' -> 'b. Costs'
    """
    # Fix lettered bullets (a, b, c, etc.) followed immediately by capital letters
    text = re.sub(r'\b([a-z])\b([A-Z])', r'\1. \2', text)
    
    # Fix numbered sub-bullets like '1Topic' -> '1. Topic'  
    text = re.sub(r'\b(\d)([A-Z])', r'\1. \2', text)
    
    # Clean up navigation symbols from ASC text
    text = re.sub(r'[>·]', '', text)
    
    # Fix spacing issues around bullets
    text = re.sub(r'([a-z])\.\s*([A-Z])', r'\1. \2', text)
    
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


def chunk_by_paragraphs(text: str, min_chunk_size: int = 200) -> List[str]:
    """
    Split text by ASC paragraph numbers (e.g., 842-10-15-1) or by double newlines
    """
    chunks = []
    
    # First try to split by ASC paragraph numbers
    paragraph_pattern = r'(\d{3}-\d{2}-\d{2}-\d+)'
    parts = re.split(paragraph_pattern, text)
    
    if len(parts) > 3:  # Successfully found paragraph numbers
        current_chunk = ""
        current_paragraph = None
        
        for i, part in enumerate(parts):
            if re.match(paragraph_pattern, part):
                # This is a paragraph number
                if current_chunk and len(current_chunk) >= min_chunk_size:
                    chunks.append(current_chunk.strip())
                current_paragraph = part
                current_chunk = f"{part}\n"
            elif current_paragraph:
                # This is content following a paragraph number
                current_chunk += part
        
        # Add the last chunk
        if current_chunk and len(current_chunk) >= min_chunk_size:
            chunks.append(current_chunk.strip())
    
    # Fallback to double newline splitting if no paragraph numbers found
    if not chunks:
        paragraphs = text.split('\n\n')
        current_chunk = ""
        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) > 1500:  # Stricter limit for embeddings
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                current_chunk += "\n\n" + paragraph if current_chunk else paragraph
        
        if current_chunk:
            chunks.append(current_chunk.strip())
    
    return [chunk for chunk in chunks if len(chunk.strip()) >= 50]


def extract_docx_text(filepath: Path) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(str(filepath))
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
        
        return '\n\n'.join(text_content)
    except Exception as e:
        logger.error(f"Error extracting text from {filepath}: {e}")
        return ""


def process_asc842_authoritative_file(filepath: Path) -> List[Dict[str, Any]]:
    """Process ASC 842 authoritative text file and return chunks with metadata"""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Clean bullet formatting issues
        content = clean_bullet_formatting(content)
        
        # Create chunks based on ASC paragraph structure (smaller chunks for embedding)
        chunks = chunk_by_paragraphs(content, min_chunk_size=200)
        
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            # Extract paragraph number if present  
            paragraph_match = re.search(r'(842-\d{2}-\d{2}-\d+)', chunk)
            paragraph_number = paragraph_match.group(1) if paragraph_match else None
            
            # Determine section from paragraph number or filename
            section = "842-10"
            section_title = "General"
            topic = "General"
            
            if paragraph_number:
                parts = paragraph_number.split('-')
                section_code = f"{parts[0]}-{parts[1]}"
                subsection = parts[2]
                
                # Map subsections to topics
                if section_code == "842-10":
                    section_title = "Overall"
                    if subsection == "05":
                        topic = "Overview"
                    elif subsection == "10":
                        topic = "Objectives"  
                    elif subsection == "15":
                        topic = "Scope"
                    elif subsection == "20":
                        topic = "Definitions"
                    elif subsection == "25":
                        topic = "Classification"
                    elif subsection == "30":
                        topic = "Initial Measurement"
                    elif subsection == "35":
                        topic = "Subsequent Measurement"
                    elif subsection == "50":
                        topic = "Disclosure"
                    elif subsection == "55":
                        topic = "Implementation"
                elif section_code == "842-20":
                    section_title = "Lessee"
                    topic = "Lessee Accounting"
                
                section = section_code
            else:
                # Fallback to filename-based detection
                filename = filepath.name.lower()
                if "overall" in filename:
                    section_title = "Overall"
                    topic = "General"
                elif "lessee" in filename:
                    section = "842-20"
                    section_title = "Lessee"  
                    topic = "Lessee Accounting"
            
            # Ensure no None values in metadata
            paragraph_number = paragraph_number or "N/A"
            
            processed_chunks.append({
                'content': chunk,
                'metadata': {
                    'source_file': filepath.name,
                    'source_type': 'authoritative',
                    'standard': 'ASC 842',
                    'section': section,
                    'section_title': section_title,
                    'topic': topic,
                    'paragraph_number': paragraph_number,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_id': f"asc842_auth_{filepath.stem}_{i}"
                }
            })
        
        logger.info(f"Processed {filepath.name}: {len(processed_chunks)} chunks, topics: {set(c['metadata']['topic'] for c in processed_chunks)}")
        return processed_chunks
        
    except Exception as e:
        logger.error(f"Error processing {filepath}: {e}")
        return []


def process_ey_interpretative_file(filepath: Path) -> List[Dict[str, Any]]:
    """Process EY interpretative guidance file and return chunks with metadata"""
    try:
        # Extract text based on file type
        if filepath.suffix.lower() == '.docx':
            content = extract_docx_text(filepath)
        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        
        if not content.strip():
            logger.warning(f"No content extracted from {filepath}")
            return []
        
        # Clean the content
        content = clean_bullet_formatting(content)
        
        # Split into chunks using paragraph-aware approach (smaller chunks for embedding)
        chunks = chunk_by_paragraphs(content, min_chunk_size=300)
        
        processed_chunks = []
        for i, chunk in enumerate(chunks):
            # Determine topic based on chunk content keywords
            chunk_lower = chunk.lower()
            topic = 'General'
            
            # Classification keywords
            classification_keywords = ['classification', 'operating lease', 'finance lease', 'ownership transfer', 'purchase option', 'lease term', 'present value', 'alternative use']
            if any(keyword in chunk_lower for keyword in classification_keywords):
                topic = 'Classification'
            else:
                # Measurement keywords  
                measurement_keywords = ['measurement', 'present value', 'discount rate', 'incremental borrowing rate', 'lease liability', 'right-of-use asset', 'initial direct costs']
                if any(keyword in chunk_lower for keyword in measurement_keywords):
                    topic = 'Measurement'
                else:
                    # Implementation keywords
                    implementation_keywords = ['implementation', 'transition', 'practical expedient', 'portfolio approach', 'commencement date']
                    if any(keyword in chunk_lower for keyword in implementation_keywords):
                        topic = 'Implementation'
                    else:
                        # Scope and identification keywords
                        scope_keywords = ['scope', 'identified asset', 'lease identification', 'control', 'substantive substitution', 'economic benefits']
                        if any(keyword in chunk_lower for keyword in scope_keywords):
                            topic = 'Identification'
                        else:
                            # Disclosure keywords
                            disclosure_keywords = ['disclosure', 'financial statement', 'reporting', 'maturity analysis', 'reconciliation']
                            if any(keyword in chunk_lower for keyword in disclosure_keywords):
                                topic = 'Disclosure'
            
            processed_chunks.append({
                'content': chunk,
                'metadata': {
                    'source_file': filepath.name,
                    'source_type': 'interpretative',
                    'standard': 'ASC 842',
                    'firm': 'EY',
                    'topic': topic,
                    'chunk_index': i,
                    'total_chunks': len(chunks),
                    'chunk_id': f"asc842_ey_{filepath.stem}_{i}"
                }
            })
        
        # Log topic distribution
        topic_counts = {}
        for chunk in processed_chunks:
            topic = chunk['metadata']['topic']
            topic_counts[topic] = topic_counts.get(topic, 0) + 1
        
        logger.info(f"Processed {filepath.name}: {len(processed_chunks)} chunks")
        logger.info(f"Topic distribution: {topic_counts}")
        return processed_chunks
        
    except Exception as e:
        logger.error(f"Error processing {filepath}: {e}")
        return []


def create_asc842_knowledge_base():
    """Create and populate ASC 842 knowledge base"""
    
    # Initialize ChromaDB client
    client = chromadb.PersistentClient(path="./asc842_knowledge_base")
    
    # Set up embedding function
    openai_ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=os.environ.get("OPENAI_API_KEY"),
        model_name="text-embedding-ada-002"
    )
    
    # Create or get collection
    collection_name = "asc842_leases"
    try:
        # Try to delete existing collection
        client.delete_collection(name=collection_name)
        logger.info(f"Deleted existing collection: {collection_name}")
    except:
        pass
    
    collection = client.create_collection(
        name=collection_name,
        embedding_function=openai_ef,
        metadata={"description": "ASC 842 Leases - Authoritative and Interpretative Guidance"}
    )
    
    # Process files from attached_assets
    attached_assets = Path("attached_assets")
    all_chunks = []
    
    # Process authoritative ASC 842 files (to be added)
    asc842_files = list(attached_assets.glob("*asc842*"))
    asc842_files.extend(list(attached_assets.glob("*842*")))
    
    for filepath in asc842_files:
        if filepath.is_file() and filepath.suffix.lower() in ['.txt', '.docx']:
            if 'ey' in filepath.name.lower():
                chunks = process_ey_interpretative_file(filepath)
            else:
                chunks = process_asc842_authoritative_file(filepath)
            all_chunks.extend(chunks)
    
    if not all_chunks:
        logger.warning("No ASC 842 files found in attached_assets. Please add:")
        logger.warning("- ASC 842 authoritative text files")
        logger.warning("- EY ASC 842 interpretative guidance")
        return
    
    # Prepare data for ChromaDB
    documents = [chunk['content'] for chunk in all_chunks]
    metadatas = [chunk['metadata'] for chunk in all_chunks]
    ids = [chunk['metadata']['chunk_id'] for chunk in all_chunks]
    
    # Add to collection in smaller batches with length validation
    batch_size = 25  # Much smaller batches to avoid token limits
    
    # Filter out any chunks that are too long (>6000 estimated tokens)
    valid_documents = []
    valid_metadatas = []
    valid_ids = []
    
    for doc, meta, doc_id in zip(documents, metadatas, ids):
        estimated_tokens = len(doc.split()) * 1.3  # Conservative estimate
        if estimated_tokens <= 6000:
            valid_documents.append(doc)
            valid_metadatas.append(meta)
            valid_ids.append(doc_id)
        else:
            logger.warning(f"Skipping oversized chunk {doc_id} ({estimated_tokens:.0f} estimated tokens)")
    
    for i in range(0, len(valid_documents), batch_size):
        batch_docs = valid_documents[i:i+batch_size]
        batch_metadata = valid_metadatas[i:i+batch_size]
        batch_ids = valid_ids[i:i+batch_size]
        
        collection.add(
            documents=batch_docs,
            metadatas=batch_metadata,
            ids=batch_ids
        )
        logger.info(f"Added batch {i//batch_size + 1}: {len(batch_docs)} chunks")
    
    # Summary statistics
    authoritative_count = len([c for c in all_chunks if c['metadata']['source_type'] == 'authoritative'])
    interpretative_count = len([c for c in all_chunks if c['metadata']['source_type'] == 'interpretative'])
    
    logger.info(f"Successfully created ASC 842 knowledge base:")
    logger.info(f"  Total chunks: {len(all_chunks)}")
    logger.info(f"  Authoritative: {authoritative_count}")
    logger.info(f"  Interpretative: {interpretative_count}")
    
    # Verify collection
    collection_count = collection.count()
    logger.info(f"Collection verification: {collection_count} documents in ChromaDB")
    
    return collection


if __name__ == "__main__":
    logger.info("Starting ASC 842 knowledge base seeding...")
    collection = create_asc842_knowledge_base()
    logger.info("ASC 842 knowledge base seeding complete!")