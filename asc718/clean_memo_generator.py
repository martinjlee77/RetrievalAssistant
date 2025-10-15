"""
ASC 718 Clean Memo Generator
"""

import streamlit as st
import streamlit.components.v1 as components
import weasyprint
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import tempfile
import os
import re
from datetime import datetime
import logging
from typing import Dict, Any, Optional
from shared.disclaimer_generator import DisclaimerGenerator

logger = logging.getLogger(__name__)

class CleanMemoGenerator:
    """Clean memo generator for ASC 718 analysis with enhanced formatting and download options."""
    
    def __init__(self, template_path=None):
        """Initialize the clean memo generator."""
        pass
    
    def combine_clean_steps(self, analysis_results: Dict[str, Any], analysis_id: Optional[str] = None) -> str:
        """Combine clean step markdown into final memo - NO PROCESSING."""
        
        # Get basic info
        customer_name = analysis_results.get('customer_name', 'Customer')
        analysis_title = analysis_results.get('analysis_title', 'Stock Compensation Analysis')
        analysis_date = datetime.now().strftime("%B %d, %Y")
        
        # Build memo with memo ID and disclaimer at very top
        memo_lines = []
        
        # Add memo ID at the very top if provided
        if analysis_id:
            memo_lines.extend([
                f"**MEMO ID:** {analysis_id}",
                "",
            ])
        
        # Add disclaimer
        memo_lines.extend([
            DisclaimerGenerator.get_top_banner(),
            "",
            "# ASC 718 MEMORANDUM",
            "",
            f"**TO:** Chief Accounting Officer",
            f"**FROM:** Technical Accounting Team - AI", 
            f"**DATE:** {analysis_date}",
            f"**RE:** {analysis_title} - ASC 718 Stock Compensation Analysis",
            f"**DOCUMENTS REVIEWED:** {analysis_results.get('filename', 'Stock Compensation Documents')}",
            "",
            ""
        ])
        
        # Add Executive Summary
        if 'executive_summary' in analysis_results:
            memo_lines.extend([
                "## EXECUTIVE SUMMARY",
                "",
                analysis_results['executive_summary'],
                "",
                ""
            ])
        
        # Add Background
        memo_lines.extend([
            "## BACKGROUND",
            ""
        ])
        
        if 'background' in analysis_results:
            memo_lines.extend([
                analysis_results['background'],
                "", 
                ""
            ])
        else:
            memo_lines.extend([
                f"We have reviewed the stock compensation documents provided by {customer_name} to determine the appropriate accounting treatment under ASC 718. This memorandum presents our analysis following the ASC 718 methodology for share-based payment arrangements.",
                "", 
                ""
            ])
        
        # Add Analysis Section Header
        memo_lines.extend([
            "## ASC 718 ANALYSIS",
        ])
        
        # Add each step's clean markdown content - check both locations
        steps_added = 0
        for step_num in range(1, 6):
            step_key = f'step_{step_num}'
            step_data = None
            
            # Check if steps are in analysis_results directly
            if step_key in analysis_results:
                step_data = analysis_results[step_key]
            # Check if steps are in analysis_results['steps']
            elif 'steps' in analysis_results and step_key in analysis_results['steps']:
                step_data = analysis_results['steps'][step_key]
            
            if step_data and isinstance(step_data, dict) and 'markdown_content' in step_data:
                # Add clean content directly - ZERO PROCESSING
                clean_content = step_data['markdown_content']
                memo_lines.append(clean_content)
                memo_lines.append("")
                steps_added += 1
                logger.info(f"Added clean step {step_num} content ({len(clean_content)} chars)")
            else:
                logger.warning(f"Step {step_num} not found or missing markdown_content. Available keys: {list(analysis_results.keys())}")
        
        # Add Conclusion Section
        if 'conclusion' in analysis_results:
            memo_lines.extend([
                "",
                "## CONCLUSION",
                "",
                analysis_results['conclusion'],
                "",
            ])
        
        # Add footer with full disclaimer
        memo_lines.extend([
            "---",
            "",
            "**PREPARED BY:** [Analyst Name] | [Title] | [Date]",
            "**REVIEWED BY:** [Reviewer Name] | [Title] | [Date]",
            "",
            DisclaimerGenerator.get_full_disclaimer()
        ])
        
        # Join and return - NO PROCESSING
        final_memo = "\n".join(memo_lines)
        logger.info(f"Clean memo generated: {len(final_memo)} chars, {steps_added}/5 steps")
        return final_memo
    
    def display_clean_memo(self, memo_content: str, analysis_id: Optional[str] = None, filename: Optional[str] = None, customer_name: Optional[str] = None):
        """Display clean memo with enhanced formatting and download options."""
        
        # Apply HTML styling for better readability
        styled_memo = self._apply_html_styling(memo_content)
        
        # Create a container with max width for better readability
        with st.container():
            st.markdown(
                f'<div style="max-width: 800px;">{styled_memo}</div>',
                unsafe_allow_html=True
            )
        
        # Add download section
        st.markdown("---")
        with st.container(border=True):
            st.info("""**IMPORTANT:** Choose your preferred format to save this memo before navigating away. The analysis results will be lost if you leave this page without saving.""")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                # Markdown download
                md_content = self._clean_markdown_content(memo_content)
                st.download_button(
                    label="📝 Markdown",
                    data=md_content,
                    file_name=f"ASC718_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
            
            with col2:
                # PDF download
                try:
                    pdf_bytes = self._generate_pdf(memo_content)
                    if pdf_bytes:
                        st.download_button(
                            label="📄 PDF",
                            data=pdf_bytes,
                            file_name=f"ASC718_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    else:
                        st.button("📄 PDF Error", disabled=True, use_container_width=True, 
                                 help="PDF generation failed")
                except Exception as e:
                    logger.error(f"PDF generation failed: {e}")
                    st.button("📄 PDF Error", disabled=True, use_container_width=True, 
                             help="PDF generation temporarily unavailable")
            
            with col3:
                # DOCX download
                try:
                    docx_bytes = self._generate_docx(memo_content)
                    if docx_bytes:
                        st.download_button(
                            label="📋 Word",
                            data=docx_bytes,
                            file_name=f"ASC718_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    else:
                        st.button("📋 DOCX Error", disabled=True, use_container_width=True,
                                 help="Word document generation failed")
                except Exception as e:
                    logger.error(f"DOCX generation failed: {e}")
                    st.button("📋 DOCX Error", disabled=True, use_container_width=True,
                             help="Word document generation temporarily unavailable")
            
            with col4:
                # Copy to clipboard button
                copy_key = f"copy_{hash(memo_content[:100])}"
                if st.button("📋 Copy", use_container_width=True, key=copy_key, help="Copy memo to clipboard"):
                    st.info("Copy functionality requires manual selection and Ctrl+C")
            
    
    def _apply_html_styling(self, content: str) -> str:
        """Apply HTML styling to markdown content for better display."""
        
        # Convert markdown headers to HTML with styling
        content = re.sub(r'^# (.+)$', r'<h1 style="color: #1f77b4; border-bottom: 2px solid #1f77b4; padding-bottom: 10px; margin-top: 30px; margin-bottom: 20px;">\1</h1>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<h2 style="color: #ffffff; margin-top: 25px; margin-bottom: 15px; font-size: 1.3em;">\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.+)$', r'<h3 style="color: #ffffff; margin-top: 20px; margin-bottom: 10px; font-size: 1.1em;">\1</h3>', content, flags=re.MULTILINE)
        
        # Style bold text
        content = re.sub(r'\*\*(.+?)\*\*', r'<strong style="color: #ffffff;">\1</strong>', content)
        
        # Convert bullet points
        content = re.sub(r'^- (.+)$', r'<li style="margin-bottom: 5px;">\1</li>', content, flags=re.MULTILINE)
        content = re.sub(r'(<li.*?>.*?</li>\s*)+', r'<ul style="margin: 10px 0; padding-left: 20px;">\g<0></ul>', content, flags=re.DOTALL)
        
        # Add paragraph styling
        lines = content.split('\n')
        styled_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith('<') and not line.startswith('|'):
                if not any(marker in line for marker in ['<h1', '<h2', '<h3', '<li', '<ul', '</ul>']):
                    line = f'<p style="margin-bottom: 15px; line-height: 1.6;">{line}</p>'
            styled_lines.append(line)
        
        return '\n'.join(styled_lines)
    
    def _clean_markdown_content(self, content: str) -> str:
        """Clean and format markdown content for download."""
        
        # Remove any HTML tags that might have been added
        content = re.sub(r'<[^>]+>', '', content)
        
        # Clean up extra whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Ensure proper markdown formatting
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1]:  # Only add blank line if previous line wasn't blank
                cleaned_lines.append('')
        
        return '\n'.join(cleaned_lines)
    
    def _generate_pdf(self, content: str) -> Optional[bytes]:
        """Generate PDF using ReportLab"""
        try:
            from shared.pdf_generator import generate_pdf_from_markdown
            
            logger.info("Generating PDF using ReportLab")
            pdf_bytes = generate_pdf_from_markdown(content)
            
            if pdf_bytes:
                logger.info(f"PDF generation successful: {len(pdf_bytes)} bytes")
                return pdf_bytes
            else:
                logger.error("PDF generation returned None")
                return None
        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            return None
    
    def _generate_docx(self, content: str) -> Optional[bytes]:
        """Generate DOCX from markdown content."""
        try:
            doc = Document()

            # Page set up
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)
            
            # Set default font for the whole document
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(6)
                        
            # Process content line by line
            lines = content.split('\n')

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Headers
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.runs[0].font.size = Pt(24)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # Better section spacing for main title
                    heading.paragraph_format.space_before = Pt(18)
                    heading.paragraph_format.space_after = Pt(12)
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                    heading.runs[0].font.size = Pt(18)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # Better section spacing for section headers
                    heading.paragraph_format.space_before = Pt(15)
                    heading.paragraph_format.space_after = Pt(9)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                    heading.runs[0].font.size = Pt(14)
                    heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
                    # Better section spacing for step headers
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(6)
                # Bold text
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line[2:-2]).bold = True
                    # Add line spacing for bold paragraphs
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                # Bullet points
                elif line.startswith('- '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                    # Add line spacing for bullet points
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                # Regular paragraphs
                else:
                    # Handle bold text
                    paragraph = doc.add_paragraph()
                    self._add_formatted_text(paragraph, line)
                    # Add line spacing for regular paragraphs  
                    paragraph.paragraph_format.space_after = Pt(6)
                    paragraph.paragraph_format.line_spacing = 1.15
                
            # Save to bytes
            with tempfile.NamedTemporaryFile() as tmp_file:
                doc.save(tmp_file.name)
                tmp_file.seek(0)
                return tmp_file.read()

        except Exception as e:
            logger.error(f"DOCX generation failed: {e}")
            return None
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to paragraph, handling bold markup."""
        
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Regular text
                paragraph.add_run(part)
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown content to HTML."""
        
        # Simple markdown to HTML conversion
        html_content = content
        
        # Headers
        html_content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_content, flags=re.MULTILINE)
        
        # Bold text
        html_content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_content)
        
        # Lists
        html_content = re.sub(r'^- (.+)$', r'<li>\1</li>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'(<li>.*</li>\s*)+', r'<ul>\g<0></ul>', html_content, flags=re.DOTALL)
        
        # Paragraphs
        lines = html_content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('<'):
                line = f'<p>{line}</p>'
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)
    
    def _convert_markdown_to_html(self, content: str) -> str:
        """Convert markdown content to HTML for display."""
        
        # Convert headers
        content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        content = re.sub(r'^#### (.+)$', r'<h4>\1</h4>', content, flags=re.MULTILINE)
        content = re.sub(r'^##### (.+)$', r'<h5>\1</h5>', content, flags=re.MULTILINE)
        content = re.sub(r'^###### (.+)$', r'<h6>\1</h6>', content, flags=re.MULTILINE)
        
        # Convert bold text
        content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', content)
        
        # Convert italic text
        content = re.sub(r'\*(.+?)\*', r'<em>\1</em>', content)
        
        # Convert bullet points
        content = re.sub(r'^- (.+)$', r'<li>\1</li>', content, flags=re.MULTILINE)
        
        # Wrap consecutive list items in <ul> tags
        content = re.sub(r'(<li>.*</li>\s*)+', r'<ul>\g<0></ul>', content, flags=re.DOTALL)
        
        # Convert line breaks to paragraphs
        paragraphs = content.split('\n\n')
        html_paragraphs = []
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if paragraph:
                # Don't wrap headers or lists in paragraph tags
                if not (paragraph.startswith('<h') or paragraph.startswith('<ul>') or paragraph.startswith('<li>')):
                    # Handle multi-line paragraphs
                    paragraph = paragraph.replace('\n', '<br>')
                    paragraph = f'<p>{paragraph}</p>'
                html_paragraphs.append(paragraph)
        
        return '\n'.join(html_paragraphs)