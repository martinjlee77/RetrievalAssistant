"""
Clean Memo Generator - No Text Processing
Displays clean GPT-4o output without any corruption.
"""

import streamlit as st
import streamlit.components.v1 as components
import logging
from datetime import datetime
from typing import Dict, Any, Optional
from shared.disclaimer_generator import DisclaimerGenerator

logger = logging.getLogger(__name__)

class CleanMemoGenerator:
    """Clean memo generator that preserves GPT-4o text exactly as generated."""
    
    def __init__(self, template_path=None):
        """Initialize - template_path ignored for now."""
        pass
    
    def combine_clean_steps(self, analysis_results: Dict[str, Any]) -> str:
        """Combine clean step markdown into final memo - NO PROCESSING."""
        
        # Get basic info
        customer_name = analysis_results.get('customer_name', 'Entity')
        analysis_title = analysis_results.get('analysis_title', 'Lease Contract Analysis')
        analysis_date = datetime.now().strftime("%B %d, %Y")
        
        # Build memo with disclaimer at very top
        memo_lines = [
            DisclaimerGenerator.get_top_banner(),
            "",
            "# ASC 842 MEMORANDUM",
            "",
            f"**TO:** Chief Accounting Officer",
            f"**FROM:** Technical Accounting Team - AI", 
            f"**DATE:** {analysis_date}",
            f"**RE:** {analysis_title} - ASC 842 Lease Accounting Analysis",
            f"**DOCUMENTS REVIEWED:** {analysis_results.get('filename', 'Lease Agreement and Related Documents')}"
            "",
            ""
        ]
        
        # Add Executive Summary
        if 'executive_summary' in analysis_results:
            memo_lines.extend([
                "",
                analysis_results['executive_summary'],
                "",
                ""
            ])
        
        # Add Background
        memo_lines.extend([
            ""
        ])
        
        if 'background' in analysis_results:
            memo_lines.extend([
                analysis_results['background'],
                "", 
                ""
            ])
        else:
            memo_lines.extend([
                f"We have reviewed the lease agreement and related documents to determine the appropriate lease accounting treatment under ASC 842. This memorandum presents our analysis following the five-step ASC 842 methodology for initial lease accounting.",
                "", 
                ""
            ])
        
        # Add Analysis Section Header
        memo_lines.extend([
            "## ASC 842 ANALYSIS",
        ])
        
        # Add each step's clean markdown content - check both locations
        steps_added = 0
        for step_num in range(1, 6):
            step_key = f'step_{step_num}'
            step_data = None
            
            # Check if steps are in analysis_results directly
            if step_key in analysis_results:
                step_data = analysis_results[step_key]
            # Check if steps are in analysis_results['steps']
            elif 'steps' in analysis_results and step_key in analysis_results['steps']:
                step_data = analysis_results['steps'][step_key]
            
            if step_data and isinstance(step_data, dict) and 'markdown_content' in step_data:
                # Add clean content directly - ZERO PROCESSING
                clean_content = step_data['markdown_content']
                memo_lines.append(clean_content)
                memo_lines.append("")
                steps_added += 1
                logger.info(f"Added clean step {step_num} content ({len(clean_content)} chars)")
            else:
                logger.warning(f"Step {step_num} not found or missing markdown_content. Available keys: {list(analysis_results.keys())}")
        
        # Add Conclusion Section
        if 'conclusion' in analysis_results:
            memo_lines.extend([
                "",
                analysis_results['conclusion'],
                "",
            ])
        
        # Add footer with full disclaimer
        memo_lines.extend([
            "---",
            "",
            "**PREPARED BY:** [Analyst Name] | [Title] | [Date]",
            "**REVIEWED BY:** [Reviewer Name] | [Title] | [Date]",
            "",
            DisclaimerGenerator.get_full_disclaimer()
        ])
        
        # Join and return - NO PROCESSING
        final_memo = "\n".join(memo_lines)
        logger.info(f"Clean memo generated: {len(final_memo)} chars, {steps_added}/5 steps")
        return final_memo
    
    def display_clean_memo(self, memo_content: str) -> None:
        """Display clean memo content with enhanced download options."""
        
        # Validate memo content
        if not memo_content or memo_content.strip() == "":
            st.error("Memo content is empty. Please regenerate the analysis.")
            return
            
        # Log what we're about to display
        logger.info(f"Displaying clean memo sample: {repr(memo_content[:150])}")
        
        # Convert markdown to HTML manually to bypass Streamlit's markdown processor
        html_content = self._convert_markdown_to_html(memo_content)
        
        # Use HTML display which preserves formatting
        st.markdown(html_content, unsafe_allow_html=True)
        
        # Enhanced Download Section - AFTER memo display (for stability)
        if memo_content and len(memo_content.strip()) > 10:
            st.markdown("---")
            st.markdown("### 💾 Save Your Memo")
            st.info("**IMPORTANT:** Choose your preferred format to save this memo before navigating away.")
            
            # Create columns for download buttons
            col1, col2, col3, col4 = st.columns(4)
            
            # Generate timestamp for consistent filenames
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            base_filename = f"asc842_memo_{timestamp}"
            
            with col1:
                # Markdown download (existing)
                st.download_button(
                    label="📄 Markdown",
                    data=memo_content,
                    file_name=f"{base_filename}.md",
                    mime="text/markdown",
                    key=f"download_md_{hash(memo_content[:100])}",
                    use_container_width=True
                )
            
            with col2:
                # PDF download
                pdf_data = self._generate_pdf(memo_content)
                if pdf_data:
                    st.download_button(
                        label="📄 PDF",
                        data=pdf_data,
                        file_name=f"{base_filename}.pdf",
                        mime="application/pdf",
                        key=f"download_pdf_{hash(memo_content[:100])}",
                        use_container_width=True
                    )
                else:
                    st.button("📄 PDF", disabled=True, use_container_width=True, help="PDF generation failed")
            
            with col3:
                # DOCX download
                docx_data = self._generate_word_document(memo_content)
                if docx_data:
                    st.download_button(
                        label="📄 Word (.docx)",
                        data=docx_data,
                        file_name=f"{base_filename}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{hash(memo_content[:100])}",
                        use_container_width=True
                    )
                else:
                    st.button("📄 Word", disabled=True, use_container_width=True, help="Word generation failed")
            
            with col4:
                # One-click copy button with JavaScript
                copy_key = f"copy_{hash(memo_content[:100])}"
                if st.button("📋 Copy to Clipboard", use_container_width=True, key=copy_key):
                    # Escape content for JavaScript
                    escaped_content = memo_content.replace('`', '\\`').replace('$', '\\$').replace('\\', '\\\\').replace('"', '\\"')
                    
                    # Create JavaScript component to copy to clipboard
                    copy_js = f"""
                    <script>
                        function copyToClipboard() {{
                            const textToCopy = `{escaped_content}`;
                            navigator.clipboard.writeText(textToCopy).then(function() {{
                                alert('Memo copied to clipboard!');
                            }}).catch(function(err) {{
                                // Fallback for older browsers
                                const textArea = document.createElement("textarea");
                                textArea.value = textToCopy;
                                document.body.appendChild(textArea);
                                textArea.select();
                                document.execCommand('copy');
                                document.body.removeChild(textArea);
                                alert('Memo copied to clipboard!');
                            }});
                        }}
                        copyToClipboard();
                    </script>
                    """
                    st.components.v1.html(copy_js, height=0)
    
    def _generate_pdf(self, content: str) -> bytes:
        """Generate PDF from markdown content."""
        try:
            import weasyprint
            from markdown import markdown
            
            # Convert markdown to HTML
            html_content = f"""
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: 'Times New Roman', serif; line-height: 1.6; margin: 40px; }}
                    h1 {{ color: #2E4F99; border-bottom: 2px solid #2E4F99; }}
                    h2 {{ color: #2E4F99; margin-top: 30px; }}
                    h3 {{ color: #333; }}
                    .disclaimer {{ background-color: #f8f9fa; padding: 15px; border-left: 4px solid #dc3545; margin: 20px 0; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
            {markdown(content, extensions=['tables', 'fenced_code'])}
            </body>
            </html>
            """
            
            # Generate PDF
            pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
            return pdf_bytes or b"PDF generation failed"
            
        except ImportError:
            logger.error("WeasyPrint not available for PDF generation")
            # Fallback to simple text PDF
            return self._generate_simple_pdf(content)
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return self._generate_simple_pdf(content)
    
    def _generate_simple_pdf(self, content: str) -> bytes:
        """Generate simple PDF using FPDF as fallback."""
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', size=11)
            
            # Split content into lines and add to PDF
            lines = content.split('\n')
            for line in lines:
                # Handle long lines
                if len(line) > 80:
                    words = line.split(' ')
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            pdf.cell(0, 5, current_line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                            current_line = word + " "
                    if current_line:
                        pdf.cell(0, 5, current_line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
                else:
                    pdf.cell(0, 5, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)
            
            pdf_output = pdf.output(dest='S')
            if isinstance(pdf_output, bytes):
                return pdf_output
            elif isinstance(pdf_output, str):
                return pdf_output.encode('latin-1')
            else:
                return bytes(pdf_output)
            
        except Exception as e:
            logger.error(f"Simple PDF generation error: {e}")
            return b"PDF generation failed"
    
    def _generate_word_document(self, content: str) -> bytes:
        """Generate Word document from markdown content."""
        try:
            from docx import Document
            from docx.shared import Inches
            import io
            
            doc = Document()
            
            # Split content into lines and process
            lines = content.split('\n')
            current_paragraph = doc.add_paragraph()
            
            for line in lines:
                line = line.strip()
                
                if line.startswith('# '):
                    # Main heading
                    heading = doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Section heading
                    heading = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Subsection heading
                    heading = doc.add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line:
                    # Regular paragraph
                    doc.add_paragraph(line)
                else:
                    # Empty line - add space
                    doc.add_paragraph()
            
            # Save to bytes
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io.read()
            
        except Exception as e:
            logger.error(f"Word document generation error: {e}")
            return b"Word document generation failed"
    
    def _convert_markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML manually to preserve currency formatting."""
        
        # Split into lines and process each one
        lines = markdown_content.split('\n')
        html_lines = []
        in_list = False
        
        for line in lines:
            # Convert headers with better spacing
            if line.startswith('# '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h1 style="margin: 20px 0 15px 0; font-weight: bold;">{line[2:]}</h1>')
            elif line.startswith('## '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h2 style="margin: 18px 0 12px 0; border-bottom: 2px solid #e0e0e0; padding-bottom: 5px;">{line[3:]}</h2>')
            elif line.startswith('### '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<h3 style="margin: 16px 0 10px 0; font-weight: 600;">{line[4:]}</h3>')
            # Handle bullet points better
            elif line.strip().startswith('- '):
                if not in_list:
                    html_lines.append('<ul style="margin: 10px 0; padding-left: 25px;">')
                    in_list = True
                content = line.strip()[2:].strip()
                # Handle bold text in list items
                if '**' in content:
                    while '**' in content:
                        content = content.replace('**', '<strong>', 1).replace('**', '</strong>', 1)
                html_lines.append(f'<li style="margin: 6px 0; line-height: 1.4;">{content}</li>')
            # Convert bold text in paragraphs
            elif '**' in line and line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                # Handle multiple bold sections
                processed_line = line
                while '**' in processed_line:
                    processed_line = processed_line.replace('**', '<strong>', 1).replace('**', '</strong>', 1)
                html_lines.append(f'<p style="margin: 10px 0; line-height: 1.6;">{processed_line}</p>')
            # Skip horizontal rules (as requested)
            elif line.strip() == '---':
                continue
            # Empty lines - add small spacing
            elif line.strip() == '':
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                continue
            # Regular paragraphs
            elif line.strip():
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                html_lines.append(f'<p style="margin: 10px 0; line-height: 1.6;">{line}</p>')
        
        # Close any open lists
        if in_list:
            html_lines.append('</ul>')
        
        # Join with improved styling
        html_content = f"""
        <div style="font-family: 'Segoe UI', 'Helvetica Neue', Arial, sans-serif; 
        line-height: 1.6; max-width: 850px; padding: 30px; 
        border: 1px solid #e1e5e9; 
        border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            {''.join(html_lines)}
        </div>
        """
        
        return html_content