#!/usr/bin/env python3
"""
Analyze the cleaned Word document for RAG suitability
"""
from docx import Document
import re

def analyze_cleaned_document(file_path):
    """Analyze the cleaned Word document"""
    
    doc = Document(file_path)
    
    analysis = {
        'overview': {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'estimated_pages': len(doc.paragraphs) // 30  # Rough estimate
        },
        'content_quality': {},
        'structure_analysis': {},
        'rag_readiness': {}
    }
    
    # Analyze content structure
    headings = []
    substantive_content = []
    tables_summary = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        # Track headings
        if para.style.name.startswith('Heading'):
            headings.append({
                'level': para.style.name,
                'text': text,
                'paragraph_index': i
            })
        
        # Collect substantive content (>100 chars, not just headers)
        if len(text) > 100 and not para.style.name.startswith('Heading'):
            substantive_content.append({
                'index': i,
                'length': len(text),
                'text': text[:150] + '...' if len(text) > 150 else text
            })
    
    # Analyze tables
    for i, table in enumerate(doc.tables):
        if table.rows:
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            tables_summary.append({
                'index': i,
                'rows': len(table.rows),
                'columns': len(table.rows[0].cells),
                'first_row': first_row
            })
    
    # Content quality assessment
    analysis['content_quality'] = {
        'total_headings': len(headings),
        'substantive_paragraphs': len(substantive_content),
        'avg_paragraph_length': sum(p['length'] for p in substantive_content) / len(substantive_content) if substantive_content else 0,
        'tables_count': len(tables_summary)
    }
    
    # Structure analysis
    analysis['structure_analysis'] = {
        'heading_levels': list(set(h['level'] for h in headings)),
        'major_sections': [h['text'] for h in headings if h['level'] == 'Heading 1'][:10],
        'step_structure': [h['text'] for h in headings if 'Step' in h['text']]
    }
    
    # Check for key ASC 606 content
    key_topics = [
        'performance obligation', 'transaction price', 'revenue recognition',
        'customer contract', 'distinct goods', 'principal agent',
        'variable consideration', 'contract modification'
    ]
    
    topic_coverage = {}
    for topic in key_topics:
        count = 0
        for para in substantive_content:
            if topic.lower() in para['text'].lower():
                count += 1
        topic_coverage[topic] = count
    
    analysis['rag_readiness'] = {
        'topic_coverage': topic_coverage,
        'content_density': len(substantive_content) / analysis['overview']['total_paragraphs'],
        'table_richness': len(tables_summary) / analysis['overview']['total_paragraphs']
    }
    
    return analysis, headings, substantive_content, tables_summary

def print_analysis(analysis, headings, substantive_content, tables_summary):
    """Print comprehensive analysis"""
    
    print("=== CLEANED WORD DOCUMENT ANALYSIS ===\n")
    
    # Overview
    print("📊 DOCUMENT OVERVIEW:")
    overview = analysis['overview']
    print(f"  • Total paragraphs: {overview['total_paragraphs']}")
    print(f"  • Total tables: {overview['total_tables']}")
    print(f"  • Estimated pages: {overview['estimated_pages']}")
    print()
    
    # Content quality
    print("📝 CONTENT QUALITY:")
    quality = analysis['content_quality']
    print(f"  • Headings: {quality['total_headings']}")
    print(f"  • Substantive paragraphs: {quality['substantive_paragraphs']}")
    print(f"  • Average paragraph length: {quality['avg_paragraph_length']:.0f} chars")
    print(f"  • Tables: {quality['tables_count']}")
    print()
    
    # Structure
    print("🏗️ STRUCTURE ANALYSIS:")
    structure = analysis['structure_analysis']
    print(f"  • Heading levels: {structure['heading_levels']}")
    print("  • Major sections:")
    for section in structure['major_sections']:
        print(f"    - {section}")
    print("  • Step structure:")
    for step in structure['step_structure']:
        print(f"    - {step}")
    print()
    
    # RAG readiness
    print("🎯 RAG READINESS ASSESSMENT:")
    rag = analysis['rag_readiness']
    print("  • Topic coverage:")
    for topic, count in rag['topic_coverage'].items():
        print(f"    - {topic}: {count} mentions")
    print(f"  • Content density: {rag['content_density']:.2f}")
    print(f"  • Table richness: {rag['table_richness']:.3f}")
    print()
    
    # Sample content
    print("📄 CONTENT SAMPLES:")
    for i, sample in enumerate(substantive_content[:5]):
        print(f"  [{sample['index']}] {sample['text']}")
        print()
    
    # Sample tables
    print("📊 TABLE SAMPLES:")
    for i, table in enumerate(tables_summary[:5]):
        print(f"  • Table {table['index']}: {table['rows']}x{table['columns']}")
        print(f"    Headers: {table['first_row']}")
        print()
    
    # Overall assessment
    print("🎉 OVERALL RAG SUITABILITY:")
    
    score = 0
    
    # Content volume (30 points)
    if quality['substantive_paragraphs'] > 1000:
        score += 30
        print("  ✅ Excellent content volume (30/30)")
    elif quality['substantive_paragraphs'] > 500:
        score += 20
        print("  ✅ Good content volume (20/30)")
    else:
        score += 10
        print("  ⚠️ Moderate content volume (10/30)")
    
    # Structure quality (25 points)
    if quality['total_headings'] > 50:
        score += 25
        print("  ✅ Excellent structure (25/25)")
    elif quality['total_headings'] > 20:
        score += 15
        print("  ✅ Good structure (15/25)")
    else:
        score += 5
        print("  ⚠️ Basic structure (5/25)")
    
    # Table richness (20 points)
    if quality['tables_count'] > 50:
        score += 20
        print("  ✅ Rich table content (20/20)")
    elif quality['tables_count'] > 20:
        score += 15
        print("  ✅ Good table content (15/20)")
    else:
        score += 5
        print("  ⚠️ Limited table content (5/20)")
    
    # Topic coverage (25 points)
    covered_topics = sum(1 for count in rag['topic_coverage'].values() if count > 0)
    if covered_topics >= 7:
        score += 25
        print("  ✅ Comprehensive topic coverage (25/25)")
    elif covered_topics >= 5:
        score += 15
        print("  ✅ Good topic coverage (15/25)")
    else:
        score += 5
        print("  ⚠️ Limited topic coverage (5/25)")
    
    print(f"\n🏆 FINAL SCORE: {score}/100")
    
    if score >= 85:
        print("🎉 EXCELLENT - Ready for RAG implementation")
    elif score >= 70:
        print("✅ GOOD - Suitable for RAG with minor adjustments")
    elif score >= 55:
        print("⚠️ MODERATE - Needs some improvements")
    else:
        print("❌ POOR - Significant work needed")

if __name__ == "__main__":
    file_path = "attached_assets/ey-frdbb3043-09-24-2024_revised4RAG_1752089042429.docx"
    
    print("Analyzing cleaned Word document...")
    analysis, headings, substantive_content, tables_summary = analyze_cleaned_document(file_path)
    print_analysis(analysis, headings, substantive_content, tables_summary)