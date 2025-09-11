"""
ASC 805 Clean Memo Generator
"""

import streamlit as st
import weasyprint
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import os
import re
import logging
from datetime import datetime
from typing import Dict, Any
from shared.disclaimer_generator import DisclaimerGenerator

logger = logging.getLogger(__name__)

class CleanMemoGenerator:
    """Clean memo generator that preserves GPT-4o text exactly as generated."""
    
    def __init__(self, template_path=None):
        """Initialize - template_path ignored for now."""
        pass
    
    def combine_clean_steps(self, analysis_results: Dict[str, Any], analysis_id: str = None) -> str:
        """Combine clean step markdown into final memo - NO PROCESSING."""
        
        # Get basic info
        customer_name = analysis_results.get('customer_name', 'Customer')
        analysis_title = analysis_results.get('analysis_title', 'Contract Analysis')
        analysis_date = datetime.now().strftime("%B %d, %Y")
        
        # Build memo with memo ID and disclaimer at very top
        memo_lines = []
        
        # Add memo ID at the very top if provided
        if analysis_id:
            memo_lines.extend([
                f"**MEMO ID:** {analysis_id}",
                "",
            ])
        
        # Add disclaimer
        memo_lines.extend([
            DisclaimerGenerator.get_top_banner(),
            "",
            "# ASC 805 MEMORANDUM",
            "",
            f"**TO:** Chief Accounting Officer",
            f"**FROM:** Technical Accounting Team - AI", 
            f"**DATE:** {analysis_date}",
            f"**RE:** {analysis_title} - ASC 805 Business Combination Analysis",
            f"**DOCUMENTS REVIEWED:** {analysis_results.get('filename', 'Transaction Documents')}",
            "",
            ""
        ])
        
        # Add Executive Summary
        if 'executive_summary' in analysis_results:
            memo_lines.extend([
                "## EXECUTIVE SUMMARY",
                "",
                analysis_results['executive_summary'],
                "",
                ""
            ])
        
        # Add Background
        memo_lines.extend([
            "## BACKGROUND",
            ""
        ])
        
        if 'background' in analysis_results:
            memo_lines.extend([
                analysis_results['background'],
                "", 
                ""
            ])
        else:
            memo_lines.extend([
                f"We have reviewed the transaction documents provided by {customer_name} to determine the appropriate accounting treatment for this business combination under ASC 805. This memorandum presents our analysis following the ASC 805 methodology for business combinations.",
                "", 
                ""
            ])
        
        # Add Analysis Section Header
        memo_lines.extend([
            "## ASC 805 ANALYSIS",
        ])
        
        # Add each step's clean markdown content - check both locations
        steps_added = 0
        for step_num in range(1, 6):
            step_key = f'step_{step_num}'
            step_data = None
            
            # Check if steps are in analysis_results directly
            if step_key in analysis_results:
                step_data = analysis_results[step_key]
            # Check if steps are in analysis_results['steps']
            elif 'steps' in analysis_results and step_key in analysis_results['steps']:
                step_data = analysis_results['steps'][step_key]
            
            if step_data and isinstance(step_data, dict) and 'markdown_content' in step_data:
                # Add clean content directly - ZERO PROCESSING
                clean_content = step_data['markdown_content']
                memo_lines.append(clean_content)
                memo_lines.append("")
                steps_added += 1
                logger.info(f"Added clean step {step_num} content ({len(clean_content)} chars)")
            else:
                logger.warning(f"Step {step_num} not found or missing markdown_content. Available keys: {list(analysis_results.keys())}")
        
        # Add Conclusion Section
        if 'conclusion' in analysis_results:
            memo_lines.extend([
                "",
                "## CONCLUSION",
                "",
                analysis_results['conclusion'],
                "",
            ])
        
        # Add footer with full disclaimer
        memo_lines.extend([
            "---",
            "",
            "**PREPARED BY:** [Analyst Name] | [Title] | [Date]",
            "**REVIEWED BY:** [Reviewer Name] | [Title] | [Date]",
            "",
            DisclaimerGenerator.get_full_disclaimer()
        ])
        
        # Join and return - NO PROCESSING
        final_memo = "\n".join(memo_lines)
        logger.info(f"Clean memo generated: {len(final_memo)} chars, {steps_added}/5 steps")
        return final_memo
    
    def display_clean_memo(self, memo_content: str, analysis_id: str = None, filename: str = None, customer_name: str = None):
        """Display clean memo with enhanced formatting and download options."""
        
        # Apply HTML styling for better readability
        styled_memo = self._apply_html_styling(memo_content)
        
        # Create a container with max width for better readability
        with st.container():
            st.markdown(
                f'<div style="max-width: 800px;">{styled_memo}</div>',
                unsafe_allow_html=True
            )
        
        # Add download section
        st.markdown("---")
        with st.container(border=True):
            st.info("""**IMPORTANT:** Choose your preferred format to save this memo before navigating away. The analysis results will be lost if you leave this page without saving.""")
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                # Markdown download
                md_content = self._clean_markdown_content(memo_content)
                st.download_button(
                    label="📝 Markdown",
                    data=md_content,
                    file_name=f"ASC805_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
            
            with col2:
                # PDF download
                try:
                    pdf_bytes = self._generate_pdf(memo_content)
                    st.download_button(
                        label="📄 PDF",
                        data=pdf_bytes,
                        file_name=f"ASC805_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as e:
                    logger.error(f"PDF generation failed: {e}")
                    st.button("📄 PDF Error", disabled=True, use_container_width=True, 
                             help="PDF generation temporarily unavailable")
            
            with col3:
                # DOCX download
                try:
                    docx_bytes = self._generate_docx(memo_content)
                    st.download_button(
                        label="📋 Word",
                        data=docx_bytes,
                        file_name=f"ASC805_Analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    logger.error(f"DOCX generation failed: {e}")
                    st.button("📋 DOCX Error", disabled=True, use_container_width=True,
                             help="Word document generation temporarily unavailable")
            
            with col4:
                # Copy to clipboard button
                copy_key = f"copy_{hash(memo_content[:100])}"
                if st.button("📋 Copy", use_container_width=True, key=copy_key, help="Copy memo to clipboard"):
                    st.info("Copy functionality requires manual selection and Ctrl+C")
            
            # Add audit pack download if analysis_id provided
            if analysis_id:
                st.markdown("---")
                st.markdown("### 📋 Audit Pack")
                from shared.audit_pack_generator import AuditPackGenerator
                audit_generator = AuditPackGenerator()
                audit_generator.add_audit_pack_download(memo_content, analysis_id, filename, customer_name)
    
    
    def _apply_html_styling(self, content: str) -> str:
        """Apply HTML styling to markdown content for better display."""
        
        # Convert markdown headers to HTML with styling
        content = re.sub(r'^# (.+)$', r'<h1 style="color: #1f77b4; border-bottom: 2px solid #1f77b4; padding-bottom: 10px; margin-top: 30px; margin-bottom: 20px;">\1</h1>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'<h2 style="color: #333; margin-top: 25px; margin-bottom: 15px; font-size: 1.3em;">\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.+)$', r'<h3 style="color: #555; margin-top: 20px; margin-bottom: 10px; font-size: 1.1em;">\1</h3>', content, flags=re.MULTILINE)
        
        # Style bold text
        content = re.sub(r'\*\*(.+?)\*\*', r'<strong style="color: #2c3e50;">\1</strong>', content)
        
        # Convert bullet points
        content = re.sub(r'^- (.+)$', r'<li style="margin-bottom: 5px;">\1</li>', content, flags=re.MULTILINE)
        content = re.sub(r'(<li.*?>.*?</li>\s*)+', r'<ul style="margin: 10px 0; padding-left: 20px;">\g<0></ul>', content, flags=re.DOTALL)
        
        # Add paragraph styling
        lines = content.split('\n')
        styled_lines = []
        for line in lines:
            line = line.strip()
            if line and not line.startswith('<') and not line.startswith('|'):
                if not any(marker in line for marker in ['<h1', '<h2', '<h3', '<li', '<ul', '</ul>']):
                    line = f'<p style="margin-bottom: 15px; line-height: 1.6;">{line}</p>'
            styled_lines.append(line)
        
        return '\n'.join(styled_lines)
    
    def _clean_markdown_content(self, content: str) -> str:
        """Clean and format markdown content for download."""
        
        # Remove any HTML tags that might have been added
        content = re.sub(r'<[^>]+>', '', content)
        
        # Clean up extra whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Ensure proper markdown formatting
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1]:  # Only add blank line if previous line wasn't blank
                cleaned_lines.append('')
        
        return '\n'.join(cleaned_lines)
    
    def _generate_pdf(self, content: str) -> bytes:
        """Generate PDF from markdown content."""
        
        # Convert markdown to HTML
        html_content = self._markdown_to_html(content)
        
        # Add CSS styling
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    margin: 0.5in;
                    color: #333;
                }}
                h1 {{
                    color: #1f77b4;
                    border-bottom: 2px solid #1f77b4;
                    padding-bottom: 10px;
                    margin-top: 30px;
                    margin-bottom: 20px;
                }}
                h2 {{
                    color: #333;
                    margin-top: 25px;
                    margin-bottom: 15px;
                }}
                h3 {{
                    color: #555;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }}
                p {{
                    margin-bottom: 15px;
                }}
                ul {{
                    margin: 10px 0;
                    padding-left: 20px;
                }}
                li {{
                    margin-bottom: 5px;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                }}
                .header {{
                    text-align: center;
                    margin-bottom: 30px;
                }}
                .date {{
                    text-align: right;
                    color: #666;
                    font-size: 0.9em;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>ASC 805 Business Combination Analysis</h1>
                <div class="date">Generated on {datetime.now().strftime("%B %d, %Y")}</div>
            </div>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF using WeasyPrint
        pdf_bytes = weasyprint.HTML(string=styled_html).write_pdf()
        return pdf_bytes
    
    def _generate_docx(self, content: str) -> bytes:
        """Generate DOCX from markdown content."""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('ASC 805 Business Combination Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_paragraph = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Add spacing
        
        # Process content line by line
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()  # Add blank line
                continue
            
            # Headers
            if line.startswith('# '):
                heading = doc.add_heading(line[2:], 1)
            elif line.startswith('## '):
                heading = doc.add_heading(line[3:], 2)
            elif line.startswith('### '):
                heading = doc.add_heading(line[4:], 3)
            # Bullet points
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Regular paragraphs
            else:
                # Handle bold text
                paragraph = doc.add_paragraph()
                self._add_formatted_text(paragraph, line)
        
        # Save to bytes
        with tempfile.NamedTemporaryFile() as tmp_file:
            doc.save(tmp_file.name)
            tmp_file.seek(0)
            return tmp_file.read()
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add formatted text to paragraph, handling bold markup."""
        
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                # Regular text
                paragraph.add_run(part)
    
    def _markdown_to_html(self, content: str) -> str:
        """Convert markdown content to HTML."""
        
        # Simple markdown to HTML conversion
        html_content = content
        
        # Headers
        html_content = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html_content, flags=re.MULTILINE)
        
        # Bold text
        html_content = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html_content)
        
        # Lists
        html_content = re.sub(r'^- (.+)$', r'<li>\1</li>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'(<li>.*</li>\s*)+', r'<ul>\g<0></ul>', html_content, flags=re.DOTALL)
        
        # Paragraphs
        lines = html_content.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if line and not line.startswith('<'):
                line = f'<p>{line}</p>'
            processed_lines.append(line)
        
        return '\n'.join(processed_lines)