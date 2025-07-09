#!/usr/bin/env python3
import os
from docx import Document

def quick_analysis(file_path):
    try:
        doc = Document(file_path)
        
        print("=== QUICK WORD DOCUMENT ANALYSIS ===\n")
        
        # Basic stats
        print("📊 DOCUMENT OVERVIEW:")
        print(f"  • Total paragraphs: {len(doc.paragraphs)}")
        print(f"  • Total tables: {len(doc.tables)}")
        print()
        
        # First 20 paragraphs to understand structure
        print("📋 FIRST 20 PARAGRAPHS:")
        for i, para in enumerate(doc.paragraphs[:20]):
            text = para.text.strip()
            if text:
                print(f"  [{i}] {para.style.name}: {text[:100]}...")
        print()
        
        # Tables summary
        print("📊 TABLES SUMMARY:")
        for i, table in enumerate(doc.tables[:5]):
            print(f"  • Table {i}: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} columns")
            if table.rows:
                first_row = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"    Headers: {first_row}")
        print()
        
        # Sample content
        print("📄 CONTENT SAMPLES:")
        content_count = 0
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text and len(text) > 50 and content_count < 5:
                print(f"  [{i}] {text[:200]}...")
                content_count += 1
        
        return True
        
    except Exception as e:
        print(f"Error: {str(e)}")
        return False

if __name__ == "__main__":
    file_path = "attached_assets/ey-frdbb3043-09-24-2024_1752054352087.docx"
    quick_analysis(file_path)